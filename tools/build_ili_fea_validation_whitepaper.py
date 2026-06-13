from __future__ import annotations

import sys
from pathlib import Path

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))
sys.path.insert(0, str(ROOT / "tools"))

from cepa_crossing.gui_server import (  # noqa: E402
    ILI_SURROGATE_BENCHMARKS,
    ili_surrogate_predict,
    ili_surrogate_validation,
)
from build_ili_fea_whitepaper import (  # noqa: E402
    BLACK,
    DARK_BLUE,
    GOLD,
    MID_GRAY,
    NAVY,
    PALE_BLUE,
    PALE_TEAL,
    add_bullet,
    add_callout,
    add_header_footer,
    add_heading,
    add_number,
    add_table,
    add_text,
    configure_section,
    configure_styles,
    set_run_font,
)


OUTPUT_DIR = ROOT / "outputs"
CHART_DIR = OUTPUT_DIR / "ili_fea_validation_charts"
OUTPUT = OUTPUT_DIR / "Automated_ILI_to_FEA_Validation_White_Paper.docx"


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.3))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        doc_pr = drawing.find(".//wp:docPr", {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"})
        if doc_pr is not None:
            doc_pr.set("descr", caption)
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(10)
    set_run_font(caption_paragraph.add_run(caption), 9.5, MID_GRAY, italic=True)


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def chart_canvas(title: str, width=1512, height=900):
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((width // 2, 36), title, fill="#203748", font=font(30, True), anchor="ma")
    return image, draw


def draw_axes(draw, box, x_label, y_label, x_ticks, y_ticks, x_map, y_map):
    left, top, right, bottom = box
    draw.line((left, bottom, right, bottom), fill="#334155", width=3)
    draw.line((left, top, left, bottom), fill="#334155", width=3)
    for value, label in x_ticks:
        x = x_map(value)
        draw.line((x, bottom, x, bottom + 9), fill="#334155", width=2)
        draw.text((x, bottom + 18), label, fill="#334155", font=font(19), anchor="ma")
    for value, label in y_ticks:
        y = y_map(value)
        draw.line((left - 9, y, left, y), fill="#334155", width=2)
        draw.line((left, y, right, y), fill="#D9E2E8", width=1)
        draw.text((left - 16, y), label, fill="#334155", font=font(19), anchor="rm")
    draw.text(((left + right) // 2, bottom + 62), x_label, fill="#203748", font=font(22, True), anchor="ma")
    draw.text((left + 8, top + 8), y_label, fill="#203748", font=font(18, True), anchor="ls")


def make_charts() -> dict[str, Path]:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    actual = np.array([row[4] for row in ILI_SURROGATE_BENCHMARKS])
    ann = np.array([ili_surrogate_predict("ann", *row[:4]) for row in ILI_SURROGATE_BENCHMARKS])
    dnn = np.array([ili_surrogate_predict("dnn", *row[:4]) for row in ILI_SURROGATE_BENCHMARKS])

    parity = CHART_DIR / "published_fea_parity.png"
    image, draw = chart_canvas("Parity against 29 published FEA parameter cases")
    box = (150, 115, 1435, 770)
    low = min(actual.min(), ann.min(), dnn.min()) - 0.025
    high = max(actual.max(), ann.max(), dnn.max()) + 0.025
    x_map = lambda value: box[0] + (value - low) / (high - low) * (box[2] - box[0])
    y_map = lambda value: box[3] - (value - low) / (high - low) * (box[3] - box[1])
    ticks = [(value, f"{value:.1f}") for value in np.arange(0.2, 0.81, 0.1)]
    draw_axes(draw, box, "Published normalized FEA failure pressure", "Embedded surrogate prediction", ticks, ticks, x_map, y_map)
    draw.line((x_map(low), y_map(low), x_map(high), y_map(high)), fill="#64748B", width=3)
    for x_value, y_value in zip(actual, ann):
        x, y = x_map(x_value), y_map(y_value)
        draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill="#0F766E", outline="white", width=2)
    for x_value, y_value in zip(actual, dnn):
        x, y = x_map(x_value), y_map(y_value)
        draw.rectangle((x - 6, y - 6, x + 6, y + 6), fill="#BD5B18", outline="white", width=2)
    draw.text((190, 145), "ANN 4-8-1", fill="#0F766E", font=font(20, True))
    draw.text((190, 180), "DNN 4-10-6-1", fill="#BD5B18", font=font(20, True))
    image.save(parity)

    residual = CHART_DIR / "surrogate_residuals.png"
    image, draw = chart_canvas("Prediction error by published parameter case", height=1080)
    case = np.arange(1, len(actual) + 1)
    for panel_index, (values, title, color) in enumerate(
        (
            (ann, "ANN residuals", "#0F766E"),
            (dnn, "DNN residuals", "#BD5B18"),
        )
    ):
        error = (values - actual) / actual * 100
        top = 125 + panel_index * 455
        box = (125, top + 55, 1440, top + 390)
        limit = max(4.0, float(np.ceil(np.max(np.abs(error)) / 2) * 2))
        x_map = lambda value: box[0] + (value - 0.5) / len(case) * (box[2] - box[0])
        y_map = lambda value: box[3] - (value + limit) / (2 * limit) * (box[3] - box[1])
        draw.text((box[0], top), title, fill="#203748", font=font(24, True))
        y_ticks = [(-limit, f"{-limit:.0f}"), (0, "0"), (limit, f"{limit:.0f}")]
        x_ticks = [(value, str(value)) for value in range(1, 30, 4)]
        draw_axes(draw, box, "Published benchmark case" if panel_index else "", "Error (%)", x_ticks, y_ticks, x_map, y_map)
        zero = y_map(0)
        draw.line((box[0], zero, box[2], zero), fill="#475569", width=2)
        bar_width = max(6, int((box[2] - box[0]) / len(case) * 0.62))
        for index, value in enumerate(error, 1):
            x = x_map(index)
            y = y_map(value)
            draw.rectangle((x - bar_width // 2, min(y, zero), x + bar_width // 2, max(y, zero)), fill=color)
    image.save(residual)

    sweeps = CHART_DIR / "parameter_response.png"
    image, draw = chart_canvas("Published and predicted parameter-response trends", width=1800, height=720)
    groups = [
        ("Depth sensitivity", np.array([0.2, 0.4, 0.5, 0.6, 0.8]), np.array([0.74, 0.64, 0.56, 0.46, 0.26]), 0),
        ("Length sensitivity", np.array([0.2, 0.4, 0.8, 1.2, 1.8]), np.array([0.69, 0.62, 0.56, 0.53, 0.52]), 1),
        ("Axial stress sensitivity", np.array([0.2, 0.4, 0.5, 0.6, 0.7, 0.8, 0.9, 1.0]), np.array([0.58, 0.57, 0.56, 0.53, 0.46, 0.37, 0.20, 0.17]), 3),
    ]
    for panel_index, (title, x, y, dimension) in enumerate(groups):
        if dimension == 0:
            ann_line = [ili_surrogate_predict("ann", value, 0.8, 3.0, 0.5) for value in x]
            dnn_line = [ili_surrogate_predict("dnn", value, 0.8, 3.0, 0.5) for value in x]
            xlabel = "d/t"
        elif dimension == 1:
            ann_line = [ili_surrogate_predict("ann", 0.5, value, 3.0, 0.5) for value in x]
            dnn_line = [ili_surrogate_predict("dnn", 0.5, value, 3.0, 0.5) for value in x]
            xlabel = "l/D"
        else:
            ann_line = [ili_surrogate_predict("ann", 0.5, 0.8, 3.0, value) for value in x]
            dnn_line = [ili_surrogate_predict("dnn", 0.5, 0.8, 3.0, value) for value in x]
            xlabel = "Axial stress / SMYS"
        left = 85 + panel_index * 575
        box = (left, 150, left + 500, 570)
        x_low, x_high = float(x.min()), float(x.max())
        y_low, y_high = 0.1, 0.8
        x_map = lambda value, b=box, lo=x_low, hi=x_high: b[0] + (value - lo) / (hi - lo) * (b[2] - b[0])
        y_map = lambda value, b=box: b[3] - (value - y_low) / (y_high - y_low) * (b[3] - b[1])
        draw.text(((box[0] + box[2]) // 2, 105), title, fill="#203748", font=font(22, True), anchor="ma")
        x_ticks = [(float(value), f"{value:g}") for value in x]
        y_ticks = [(value, f"{value:.1f}") for value in (0.2, 0.4, 0.6, 0.8)]
        draw_axes(draw, box, xlabel, "Norm. pressure", x_ticks, y_ticks, x_map, y_map)
        for values, color, width in ((y, "#334155", 4), (ann_line, "#0F766E", 3), (dnn_line, "#BD5B18", 3)):
            points = [(x_map(float(xv)), y_map(float(yv))) for xv, yv in zip(x, values)]
            draw.line(points, fill=color, width=width, joint="curve")
        for xv, yv in zip(x, y):
            px, py = x_map(float(xv)), y_map(float(yv))
            draw.ellipse((px - 5, py - 5, px + 5, py + 5), fill="#334155")
    draw.text((90, 665), "Published FEA", fill="#334155", font=font(18, True))
    draw.text((300, 665), "ANN", fill="#0F766E", font=font(18, True))
    draw.text((405, 665), "DNN", fill="#BD5B18", font=font(18, True))
    image.save(sweeps)

    burst = CHART_DIR / "published_fea_vs_burst.png"
    specimens = ["X52 T1", "X52 T5", "X52 T6", "X80 IDTS2", "X80 IDTS3", "X80 IDTS4"]
    measured = np.array([23.2, 28.6, 28.7, 22.68, 20.31, 21.14])
    fea = np.array([22.95, 28.35, 27.00, 22.40, 20.12, 20.62])
    image, draw = chart_canvas("Independent full-scale validation reported for source FEA methodology")
    box = (140, 130, 1435, 750)
    x_map = lambda value: box[0] + (value + 0.5) / len(specimens) * (box[2] - box[0])
    y_map = lambda value: box[3] - value / 32.0 * (box[3] - box[1])
    x_ticks = [(index, label) for index, label in enumerate(specimens)]
    y_ticks = [(value, str(value)) for value in range(0, 33, 5)]
    draw_axes(draw, box, "Published test specimen", "Failure pressure (MPa)", x_ticks, y_ticks, x_map, y_map)
    group_width = (box[2] - box[0]) / len(specimens)
    bar_width = int(group_width * 0.28)
    for index, (test, prediction) in enumerate(zip(measured, fea)):
        center = x_map(index)
        draw.rectangle((center - bar_width - 3, y_map(test), center - 3, box[3]), fill="#334155")
        draw.rectangle((center + 3, y_map(prediction), center + bar_width + 3, box[3]), fill="#0F766E")
        error = (prediction - test) / test * 100
        draw.text((center, y_map(max(test, prediction)) - 16), f"{error:.2f}%", fill="#7F1D1D", font=font(17, True), anchor="ms")
    draw.text((190, 155), "Full-scale burst", fill="#334155", font=font(20, True))
    draw.text((190, 190), "Published FEA", fill="#0F766E", font=font(20, True))
    image.save(burst)

    return {"parity": parity, "residual": residual, "sweeps": sweeps, "burst": burst}


def build_document(charts: dict[str, Path]) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ann = ili_surrogate_validation("ann")
    dnn = ili_surrogate_validation("dnn")
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_header_footer(doc.sections[0])

    add_text(doc, "", after=86)
    add_text(doc, "VALIDATION WHITE PAPER", bold=True, color=GOLD, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    set_run_font(paragraph.add_run("Automated ILI-to-FEA"), 30, NAVY, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(5)
    set_run_font(paragraph.add_run("Published-data validation of ANN/DNN surrogate predictions"), 15, DARK_BLUE)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(28)
    set_run_font(paragraph.add_run("and the supporting finite-element methodology"), 15, DARK_BLUE)
    add_text(
        doc,
        "Benchmark definition, comparison plots, error analysis, applicability limits, and validation roadmap",
        italic=True,
        color=GOLD,
        after=88,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_text(doc, "Prepared for the Pipeline Engineering Assessment Software", bold=True, color=NAVY, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Version 1.0 | June 11, 2026", color=MID_GRAY, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)
    add_heading(doc, "Executive Summary", 1)
    add_text(
        doc,
        "This paper validates the rapid corrosion-capacity surrogate used by the Automated ILI-to-FEA workflow against 29 normalized finite-element cases published by Lo, Karuppanan, and Ovinis. It also reviews independent full-scale burst comparisons reported for the source finite-element methodology. Validation evidence is intentionally kept outside the operating interface.",
    )
    add_callout(
        doc,
        "Principal result",
        f"The embedded ANN reproduced the 29 published FEA cases with R-squared {ann['r_squared']:.5f} and MAPE {ann['mape_percent']:.2f}%. "
        f"The embedded DNN produced R-squared {dnn['r_squared']:.5f} and MAPE {dnn['mape_percent']:.2f}%. "
        "These are reproduction metrics on published data, not independent certification.",
    )
    add_heading(doc, "Validation Conclusions", 2)
    for item in (
        "Both surrogate architectures closely reproduce the selected published FEA parameter cases.",
        "The DNN has lower in-sample reproduction error; this does not prove superior extrapolation or field performance.",
        "The source literature reports close agreement between nonlinear FEA and full-scale burst tests, supporting the numerical methodology used to generate training data.",
        "The standards-based pressure assessment remains governing because the surrogate domain is limited and the validation dataset is not operator-specific.",
        "Production qualification still requires independent holdout data, raw-ILI geometry validation, nonlinear solver verification, and operator-approved reliability targets.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "1. Validation Claim and Boundaries", 1)
    add_text(
        doc,
        "The validation claim is deliberately narrow: the embedded neural networks reproduce published normalized FEA failure-pressure trends within the parameter ranges used for training. The claim does not extend to arbitrary defect morphologies, cracks, dents, SCC colonies, unknown material curves, or raw-tool signal inversion.",
    )
    add_table(
        doc,
        ["Layer", "Evidence", "Current claim"],
        [
            ["Surrogate reproduction", "29 published normalized FEA cases", "Quantified by parity, residuals, R-squared, MAE, and MAPE"],
            ["Source FEA methodology", "Six published full-scale burst comparisons from X52 and X80 tests", "Supports the underlying numerical approach; not a validation of this browser renderer"],
            ["ILI reconstruction", "Software unit and visual tests", "Geometry and data-flow verification only"],
            ["Operational acceptance", "B31.8/API RP 1183 screening and engineering review", "Not replaced by ML validation"],
        ],
        [2000, 3380, 3980],
    )
    add_callout(
        doc,
        "Governance",
        "Surrogate results are advisory. A prediction outside the published domain or near an operating limit must be replaced by an applicable standard method or a validated nonlinear FEA/ECA.",
        fill=PALE_BLUE,
    )

    add_page_break(doc)
    add_heading(doc, "2. Published Dataset", 1)
    add_text(
        doc,
        "The benchmark cases vary corrosion depth ratio d/t, axial defect length ratio l/D, normalized longitudinal spacing s/sqrt(Dt), and axial compressive stress ratio. The target is normalized failure pressure from the published finite-element study [1].",
    )
    add_table(
        doc,
        ["Variable", "Embedded range", "Role"],
        [
            ["d/t", "0.20 to 0.80", "Remaining ligament severity"],
            ["l/D", "0.20 to 1.80", "Axial defect extent"],
            ["s/sqrt(Dt)", "0 to 3.0", "Longitudinal interaction spacing"],
            ["Axial stress / SMYS", "0 to 1.0", "Combined-load severity"],
            ["Normalized pressure", "0.17 to 0.75", "Published FEA target"],
        ],
        [2600, 2200, 4560],
    )
    add_text(
        doc,
        "The ANN is a 4-8-1 feedforward network. The DNN is a 4-10-6-1 feedforward network. Both use hyperbolic-tangent hidden activations and a linear output. Inputs are linearly normalized over the published range.",
    )
    add_heading(doc, "Evaluation Measures", 2)
    add_text(
        doc,
        "R-squared measures variance explained; MAE measures mean absolute normalized-pressure error; MAPE expresses error relative to each published target; maximum absolute percentage error identifies the worst reproduced case.",
    )

    add_page_break(doc)
    add_heading(doc, "3. Parity Against Published FEA", 1)
    add_figure(doc, charts["parity"], "Figure 1. ANN and DNN predictions plotted against the 29 published normalized FEA failure-pressure cases.")
    add_text(
        doc,
        "Points close to the 45-degree line indicate agreement. Both models track the published cases across the assessed pressure range. The DNN points are visually closer because it has greater parameter capacity and was evaluated on the same published benchmark used for fitting.",
    )
    add_table(
        doc,
        ["Metric", "ANN 4-8-1", "DNN 4-10-6-1"],
        [
            ["Cases", str(ann["benchmark_case_count"]), str(dnn["benchmark_case_count"])],
            ["R-squared", f"{ann['r_squared']:.5f}", f"{dnn['r_squared']:.5f}"],
            ["MAE", f"{ann['mae_normalized_pressure']:.5f}", f"{dnn['mae_normalized_pressure']:.5f}"],
            ["MAPE", f"{ann['mape_percent']:.2f}%", f"{dnn['mape_percent']:.2f}%"],
            ["Maximum absolute error", f"{ann['maximum_absolute_error_percent']:.2f}%", f"{dnn['maximum_absolute_error_percent']:.2f}%"],
        ],
        [3600, 2880, 2880],
    )

    add_page_break(doc)
    add_heading(doc, "4. Residual and Parameter-Response Review", 1)
    add_figure(doc, charts["residual"], "Figure 2. Percentage residuals for each of the 29 published benchmark cases.")
    add_text(
        doc,
        "The ANN residuals are small but structured, indicating modest approximation bias in portions of the nonlinear response surface. The DNN residuals are smaller on this reproduction set. Residual magnitude alone cannot establish extrapolation performance.",
    )
    add_figure(doc, charts["sweeps"], "Figure 3. Published and predicted response trends for depth, length, and axial compressive stress.")
    add_text(
        doc,
        "The most important physical trend is retained: failure pressure falls strongly with increasing depth and axial compression, while length sensitivity becomes less pronounced for longer defects. Retaining monotonic physical trends is necessary but not sufficient for model qualification.",
    )

    add_page_break(doc)
    add_heading(doc, "5. Published Full-Scale Validation of the FEA Basis", 1)
    add_figure(doc, charts["burst"], "Figure 4. Published nonlinear FEA predictions compared with six full-scale burst tests reported in the source literature [2].")
    add_text(
        doc,
        "These six cases are not predictions from the embedded ANN/DNN. They are independent burst comparisons reported for related nonlinear FEA procedures used to generate corrosion-capacity datasets. Their purpose is to demonstrate that the numerical training-data methodology can reproduce measured burst pressure when material, geometry, loading, boundary conditions, and failure criteria are properly modeled.",
    )
    add_table(
        doc,
        ["Specimen", "Burst MPa", "Published FEA MPa", "Difference"],
        [
            ["X52 Test 1", "23.20", "22.95", "-1.08%"],
            ["X52 Test 5", "28.60", "28.35", "-0.87%"],
            ["X52 Test 6", "28.70", "27.00", "-5.92%"],
            ["X80 IDTS 2", "22.68", "22.40", "-1.23%"],
            ["X80 IDTS 3", "20.31", "20.12", "-0.94%"],
            ["X80 IDTS 4", "21.14", "20.62", "-2.46%"],
        ],
        [2700, 1900, 2460, 2300],
    )
    add_callout(
        doc,
        "Interpretation",
        "The burst comparisons support the source FEA methodology, but they do not validate every raw-ILI reconstruction, interaction rule, material model, crack representation, or dent assessment implemented in the software.",
    )

    add_page_break(doc)
    add_heading(doc, "6. Reliability and Uncertainty Implications", 1)
    add_text(
        doc,
        "A high surrogate R-squared does not imply a low probability of pipeline failure. Reliability depends on the separation between uncertain demand and uncertain capacity. The workflow therefore propagates inspection sizing, pressure, and model-form uncertainty through Monte Carlo simulation and reports probability of failure and reliability index separately from deterministic pressure margin.",
    )
    add_heading(doc, "Uncertainty Sources", 2)
    for item in (
        "ILI depth, length, circumferential extent, and feature-registration uncertainty.",
        "Material yield, tensile strength, toughness, stress-strain curve, and wall-thickness variability.",
        "Pressure history, axial/secondary stress, residual stress, and future growth.",
        "Model-form error from interpolation, defect idealization, interaction treatment, failure criterion, and surrogate approximation.",
        "Sampling error in rare-event probability estimates.",
    ):
        add_bullet(doc, item)
    add_text(
        doc,
        "The present coefficients of variation are user inputs. They must be calibrated from tool qualification records, excavations, material testing, pressure histories, and operator-specific model bias studies before reliability results are used for integrity decisions.",
    )

    add_heading(doc, "7. Mesh Validation Relevance", 1)
    add_text(
        doc,
        "The reconstructed pipe now uses a single shell mesh whose base surface cells are subdivided near each anomaly. No second mesh is superimposed. Standard refinement divides each local cell into 4 by 4 subcells, increasing local surface element density by approximately 16 times relative to remote cells. Transition cells reduce the element-size jump.",
    )
    add_text(
        doc,
        "Visual refinement is not a convergence study. A production FEA validation must demonstrate that failure pressure, local strain away from mathematical singularities, and fracture parameters stabilize as local element size decreases.",
    )

    add_page_break(doc)
    add_heading(doc, "8. Limitations", 1)
    for item in (
        "The 29-point comparison is a reproduction check, not an independent blind holdout.",
        "The DNN's very low benchmark error may reflect increased flexibility on a small dataset; it should not be preferred solely on in-sample error.",
        "The published full-scale comparisons validate related FEA procedures, not the browser visualization.",
        "The surrogate domain is corrosion-specific and does not cover cracks, SCC, dents, weld defects, or arbitrary mixed anomalies.",
        "Normalized published data cannot capture operator-specific material, inspection, construction, and loading uncertainty.",
        "B31.8 and API RP 1183 decisions require the applicable edition, complete assessment inputs, and qualified engineering judgment.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "9. Qualification Roadmap", 1)
    add_number(doc, "Freeze a versioned benchmark dataset with provenance, units, geometry definitions, and acceptance tolerances.")
    add_number(doc, "Reserve independent blind cases before fitting ANN/DNN parameters.")
    add_number(doc, "Export reconstructed meshes to a validated nonlinear solver and complete element-size convergence studies.")
    add_number(doc, "Validate raw MFL, caliper, and crack reconstructions against excavation or laser-scan geometry.")
    add_number(doc, "Add independent full-scale tests for corrosion interaction, dents, cracks, SCC colonies, and mixed defects.")
    add_number(doc, "Calibrate probabilistic distributions, correlations, and model bias for the operator's tools and population.")
    add_number(doc, "Establish formal model-change control, reviewer approval, and periodic revalidation.")

    add_heading(doc, "10. Conclusion", 1)
    add_text(
        doc,
        "The embedded ANN and DNN accurately reproduce the selected published corrosion-FEA cases, and the source literature demonstrates credible agreement between nonlinear FEA and full-scale burst tests. This supports use of the surrogates for rapid advisory screening within their domain. It does not justify replacing standards-based assessment, validated nonlinear FEA, or qualified integrity engineering.",
    )

    add_page_break(doc)
    add_heading(doc, "References", 1)
    references = [
        "[1] Lo, M., Karuppanan, S., and Ovinis, M. Failure Pressure Prediction of a Corroded Pipeline with Longitudinally Interacting Corrosion Defects Subjected to Combined Loadings Using FEM and ANN. Journal of Marine Science and Engineering 9(3), 281, 2021. https://doi.org/10.3390/jmse9030281",
        "[2] Lo, M., Karuppanan, S., and Ovinis, M. Artificial Neural Network-Based Failure Pressure Prediction of API 5L X52, X65 and X80 Corroded Pipes with Circumferentially Interacting Defects Subjected to Combined Loadings. Materials 15, 2022. https://pmc.ncbi.nlm.nih.gov/articles/PMC8953741/",
        "[3] Kiefner, J. F., Vieth, P. H., and Roytman, I. Continued Validation of RSTRENG, Final Report, 1996. https://www.osti.gov/biblio/441672",
        "[4] ASME. ASME B31.8, Gas Transmission and Distribution Piping Systems. https://www.asme.org/codes-standards/find-codes-standards/b31-8-gas-transmission-distribution-piping-systems",
        "[5] American Petroleum Institute. API Recommended Practice 1183, Assessment and Management of Dents in Pipelines. https://www.api.org/products-and-services/standards/important-standards-announcements/rp1183",
        "[6] American Petroleum Institute. Addendum 1 to API RP 1183 for Improved Dent Screening, May 2024. https://www.api.org/products-and-services/standards/important-standards-announcements/addendum-1-rp-1183",
    ]
    for reference in references:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.22)
        paragraph.paragraph_format.first_line_indent = Inches(-0.22)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.15
        set_run_font(paragraph.add_run(reference), 9.5, BLACK)

    doc.core_properties.title = "Automated ILI-to-FEA Validation White Paper"
    doc.core_properties.subject = "Published-data validation of ANN/DNN pipeline corrosion surrogates"
    doc.core_properties.author = "Pipeline Engineering Assessment Software"
    doc.core_properties.keywords = "ILI, FEA, validation, ANN, DNN, pipeline corrosion, burst pressure"
    doc.save(OUTPUT)


if __name__ == "__main__":
    charts = make_charts()
    build_document(charts)
    print(OUTPUT)
