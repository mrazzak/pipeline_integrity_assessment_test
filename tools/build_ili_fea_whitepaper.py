from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
OUTPUT = OUTPUT_DIR / "Automated_ILI_to_FEA_Methodology_White_Paper.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
TEAL = "0B7A75"
PALE_BLUE = "EAF3F8"
PALE_TEAL = "E9F5F3"
PALE_GRAY = "F4F6F9"
MID_GRAY = "6B7280"
WHITE = "FFFFFF"
BLACK = "111827"
GOLD = "B18A34"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[int], indent=120) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text(doc, text: str, *, bold=False, italic=False, color=BLACK, after=8, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, 11, color, bold, italic)
    return p


def add_rich_paragraph(doc, parts, *, after=8, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    if align is not None:
        p.alignment = align
    for text, opts in parts:
        run = p.add_run(text)
        set_run_font(
            run,
            opts.get("size", 11),
            opts.get("color", BLACK),
            opts.get("bold"),
            opts.get("italic"),
            opts.get("name", "Calibri"),
        )
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    for run in p.runs:
        set_run_font(run, 11, BLACK)
    if not p.runs:
        set_run_font(p.add_run(text), 11, BLACK)
    else:
        p.runs[0].text = text
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], 11, BLACK)
    else:
        set_run_font(p.add_run(text), 11, BLACK)
    return p


def add_equation(doc, equation: str, explanation: str | None = None):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2 if explanation else 0)
    run = p.add_run(equation)
    set_run_font(run, 11, DARK_BLUE, bold=True, name="Cambria")
    if explanation:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        set_run_font(p2.add_run(explanation), 9, MID_GRAY, italic=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label: str, text: str, fill=PALE_TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{label}: "), 10.5, TEAL, bold=True)
    set_run_font(p.add_run(text), 10.5, BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, PALE_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), 9, DARK_BLUE, bold=True)
    for row_values in rows:
        row = table.add_row()
        for index, text in enumerate(row_values):
            cell = row.cells[index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == len(row_values) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(str(text)), 9, BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for level, size, color, before, after in (
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.45)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p.add_run("AUTOMATED ILI-TO-FEA"), 8.5, MID_GRAY, bold=True)
    set_run_font(p.add_run("   |   Methodology White Paper"), 8.5, MID_GRAY)
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BLUE)
    pbdr.append(bottom)
    p_pr.append(pbdr)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    set_run_font(p.add_run("Engineering methodology and implementation basis   |   June 2026   |   "), 8, MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_section(section)
    add_header_footer(section)

    # Editorial cover.
    add_text(doc, "", after=86)
    add_text(doc, "ENGINEERING WHITE PAPER", bold=True, color=GOLD, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("Automated ILI-to-FEA"), 30, NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("Defect-aware geometry, screening, interaction,"), 15, DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_run_font(p.add_run("fatigue life, crack growth, and maximum MOP"), 15, DARK_BLUE)
    add_text(
        doc,
        "Methodology, implementation basis, verification requirements, and literature references",
        italic=True,
        color=GOLD,
        after=88,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_text(doc, "Prepared for the Pipeline Engineering Assessment Software", bold=True, color=NAVY, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Version 1.0 | June 6, 2026", color=MID_GRAY, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)
    add_heading(doc, "Executive Summary", 1)
    add_text(
        doc,
        "The Automated ILI-to-FEA workflow converts a summarized anomaly list, separate raw MFL/caliper/crack-tool files, or a combination of both into a common defect geometry. The same effective depth, length, width, remaining ligament, crack opening, and interaction spacing are then used by the screening calculations and the three-dimensional mesh representation.",
    )
    add_callout(
        doc,
        "Central principle",
        "Inspection geometry is not decoration. Metal loss changes the shell radius and remaining ligament; caliper data deform the shell; crack and SCC records create open discontinuities by removing intersected shell elements. Those geometry measures also replace the summarized feature dimensions used in pressure, fracture, fatigue, and interaction calculations.",
    )
    add_heading(doc, "What the Current Implementation Does", 2)
    for item in (
        "Imports TXT, CSV, tab-delimited, semicolon-delimited, and XLSX data for MFL, caliper, and crack tools independently.",
        "Allows the summarized ILI feature list to be omitted when raw tool files contain feature identifiers and location metadata.",
        "Constructs a defect-aware shell visualization without spherical or circular anomaly substitutes.",
        "Calculates isolated screening pressure, interacting-defect pressure, fatigue life, crack growth, and a governing maximum recommended operating pressure.",
        "Preserves traceability by reporting whether each feature geometry came from raw tool data or the summarized feature list.",
    ):
        add_bullet(doc, item)
    add_heading(doc, "What Requires a Production FEA Solver", 2)
    add_text(
        doc,
        "The browser model is a solver-ready geometry and engineering-assessment layer; it is not a substitute for a validated nonlinear finite-element run. A production Level 3 analysis must export the mesh to a solver, assign a calibrated elastic-plastic material curve, apply pressure and end loading, model crack faces or cohesive/XFEM behavior as appropriate, establish convergence, and validate against benchmark or full-scale test data [1, 2, 8, 9].",
    )

    add_page_break(doc)
    add_heading(doc, "1. Workflow Architecture", 1)
    add_text(
        doc,
        "The workflow has five linked stages. Each stage produces a controlled data object consumed by the next, so a change in raw geometry propagates into both the visual mesh and the assessment outputs.",
    )
    add_number(doc, "Ingest summarized ILI features and/or independent MFL, caliper, and crack-tool measurements.")
    add_number(doc, "Normalize units, feature identifiers, chainage, clock position, and local axial/circumferential coordinates.")
    add_number(doc, "Translate measurements into remaining-wall, deformed-shell, and crack-discontinuity geometry.")
    add_number(doc, "Run feature screening, pair interaction, fracture, fatigue, and maximum-MOP calculations using the translated geometry.")
    add_number(doc, "Generate the defect-aware 3D shell and an auditable report of assumptions, governing cases, and limitations.")
    add_heading(doc, "Geometry Source Policy", 2)
    add_table(
        doc,
        ["Mode", "Feature list", "Raw tool data", "Geometry used"],
        [
            ["Auto", "Optional", "Optional", "Raw data where present; feature dimensions otherwise"],
            ["Raw", "Optional", "Required", "Raw measurements define geometry and defect class"],
            ["Feature", "Required", "Ignored for sizing", "Summarized feature dimensions only"],
        ],
        [1200, 1550, 1550, 5060],
    )
    add_callout(
        doc,
        "Data governance",
        "Raw inspection values are measurements with tool uncertainty, not exact material boundaries. The workflow retains sizing tolerances and provides nominal, conservative, and probabilistic cases. Excavation or high-resolution scan data should supersede ILI dimensions when available.",
        fill=PALE_BLUE,
    )

    add_heading(doc, "2. Import Schemas and Normalization", 1)
    add_text(
        doc,
        "Each raw-tool uploader accepts a header row using the fields below. TXT files may use comma, tab, or semicolon separators. XLSX files are read from the first worksheet. Numeric cells are normalized to SI engineering units before geometry construction.",
    )
    add_table(
        doc,
        ["Source", "Required fields", "Optional but preferred"],
        [
            ["MFL", "feature_id, axial_offset_mm, circumferential_offset_mm, depth_percent", "distance_m, clock_position"],
            ["Crack", "feature_id, axial_offset_mm, circumferential_offset_mm, depth_mm, opening_mm", "distance_m, clock_position, orientation_deg, crack_id, anomaly_type"],
            ["Caliper", "feature_id, axial_offset_mm, circumferential_offset_mm, radial_deformation_mm", "distance_m, clock_position"],
        ],
        [1200, 4360, 3800],
    )
    add_text(
        doc,
        "If no summarized feature list is provided, the importer groups measurements by feature_id, derives defect type from the loaded tool, and builds a feature location from distance and clock metadata. Where a raw file omits location metadata, the local origin is used and the result must be treated as provisional.",
    )

    add_page_break(doc)
    add_heading(doc, "3. Defect Geometry Translation", 1)
    add_heading(doc, "3.1 Metal Loss from MFL", 2)
    add_text(
        doc,
        "MFL depth values are converted from percent wall loss to local depth. The shell outer surface is recessed while the inner surface is retained for an external-metal-loss interpretation. Local interpolation produces a continuous remaining-wall field rather than a single idealized pit.",
    )
    add_equation(doc, "d_i = (p_i / 100) t ;     t_rem,i = max(t - d_i, 0.01t)", "p_i is reported MFL depth percent and t is nominal wall thickness.")
    add_text(
        doc,
        "The browser visualization uses localized weighted interpolation around each measurement point. A production mesh should use triangulation, radial-basis interpolation, kriging, or a measured river-bottom surface with smoothing limited by tool spatial resolution. Realistic inspection-derived corrosion profiles and automated meshing are consistent with published Level 3 approaches [8, 9, 10].",
    )
    add_heading(doc, "3.2 Dent Geometry from Caliper", 2)
    add_text(
        doc,
        "Caliper radial deformation changes the local pipe radius continuously. Negative radial displacement produces inward denting; positive displacement may represent ovality or outward deformation. The shell coordinates are modified before the surface mesh is created so the mesh follows the deformed centerline and curvature.",
    )
    add_equation(doc, "R'(x, θ) = R + u_r(x, θ)", "u_r is the interpolated caliper radial deformation.")
    add_heading(doc, "3.3 Crack and SCC Discontinuities", 2)
    add_text(
        doc,
        "Crack records are ordered into paths by crack identifier and local coordinates. Shell triangles intersecting the crack path and opening envelope are removed, leaving an actual topological discontinuity. Crack lip lines show the two free edges. SCC is represented by multiple neighboring crack paths rather than by a filled ellipse.",
    )
    add_equation(doc, "remove(e) if dist(centroid(e), crack path) <= w_open / 2 + h_e / 2", "h_e is a local element-size allowance used to avoid bridging the crack.")
    add_callout(
        doc,
        "Mesh interpretation",
        "Element removal is suitable for a pre-existing open discontinuity. Detailed fracture analysis may instead require duplicate crack-face nodes, quarter-point elements, cohesive elements, VCCT, or XFEM. The selected technique must match the intended fracture parameter and solver.",
    )

    add_heading(doc, "4. Geometry Coupling to Assessment", 1)
    add_text(
        doc,
        "For every translated feature the workflow records maximum effective depth, axial and circumferential spans, minimum remaining ligament, average MFL area ratio, maximum crack opening, and raw point counts. These values replace the summarized dimensions before screening. Therefore, a deeper MFL point lowers remaining strength, a longer raw crack changes fracture and interaction response, and denser crack paths increase the number of removed mesh elements.",
    )
    add_table(
        doc,
        ["Derived quantity", "Raw basis", "Assessment use"],
        [
            ["Effective depth", "Maximum MFL loss, crack depth, or absolute caliper deformation", "Pressure capacity, fracture, severity"],
            ["Effective length/width", "Span of local raw coordinates", "Folias factor, interaction spacing, mesh footprint"],
            ["Remaining ligament", "Nominal wall minus effective depth", "Local refinement and collapse relevance"],
            ["Crack opening/path", "Crack-tool points grouped by crack_id", "Element removal and crack/SCC classification"],
        ],
        [2300, 3480, 3580],
    )

    add_page_break(doc)
    add_heading(doc, "5. Remaining-Strength Methods", 1)
    add_heading(doc, "5.1 Modified ASME B31G", 2)
    add_text(
        doc,
        "The implementation follows the Modified B31G form used for volumetric metal-loss screening. Flow stress is taken as SMYS + 69 MPa and optionally capped at SMTS. The failure pressure is then multiplied by the selected assessment factor. ASME B31G is the governing standards reference for this class of assessment [1].",
    )
    add_equation(doc, "z = L² / (Dt)")
    add_equation(doc, "M = √(1 + 0.6275z - 0.003375z²), z ≤ 50;     M = 3.3 + 0.032z, z > 50")
    add_equation(doc, "σ_f = σ_flow [1 - 0.85(d/t)] / [1 - 0.85(d/t)/M]")
    add_equation(doc, "P_f = 2tσ_f / D ;     P_allow = F_A P_f")
    add_heading(doc, "5.2 RSTRENG Effective Area", 2)
    add_text(
        doc,
        "RSTRENG uses the measured axial depth profile instead of a single idealized rectangle. The implementation evaluates every contiguous station interval, integrates the metal-loss area by the trapezoidal rule, calculates the corresponding Folias factor, and selects the interval with the lowest predicted failure pressure. RSTRENG was developed to reduce the over-conservatism of the original B31G criterion and was validated against corroded-pipe burst tests [3, 4].",
    )
    add_equation(doc, "A_d = Σ [(d_i + d_{i+1}) / 2] (x_{i+1} - x_i) ;     A_0 = Lt")
    add_equation(doc, "P_f = (2tσ_flow / D) [1 - A_d/A_0] / [1 - (A_d/A_0)/M]")
    add_heading(doc, "5.3 Method Selection", 2)
    add_text(
        doc,
        "The ranking workflow can apply a calculation method per feature. Metal loss may use Modified B31G, original B31G, simplified RSTRENG, CorLAS, or reported pressure; crack-like features may use fracture or SCC-colony methods. The governing result is the lowest allowable pressure among applicable methods, not the method with the highest apparent precision.",
    )

    add_page_break(doc)
    add_heading(doc, "6. Crack, SCC, and Fatigue Methodology", 1)
    add_heading(doc, "6.1 Fracture Screening", 2)
    add_text(
        doc,
        "For a crack of depth a, the linear-elastic screening stress intensity is calculated with a geometry factor Y. Axial cracks use hoop stress from internal pressure; circumferential cracks use longitudinal pressure stress. The simple expression is a screening relation and does not replace a full failure-assessment diagram or elastic-plastic fracture analysis under API 579-1/ASME FFS-1 [2].",
    )
    add_equation(doc, "K_I = Y σ √(πa) ;     σ_h = PD/(2t) ;     σ_L = PD/(4t)")
    add_heading(doc, "6.2 Paris-Law Crack Growth", 2)
    add_text(
        doc,
        "The crack-growth module numerically integrates the Paris-Erdogan relation from the initial to critical crack depth using midpoint increments. Growth is set to zero when the calculated stress-intensity range is at or below the selected threshold. Material C and m values must be consistent with the units used by the implementation and supported by test data [5].",
    )
    add_equation(doc, "da/dN = C(ΔK - ΔK_th)^m ;     ΔK = Y Δσ √(πa)")
    add_equation(doc, "N = ∫[a_0 to a_c] da / {C(ΔK - ΔK_th)^m}")
    add_heading(doc, "6.3 SCC Colony Interaction", 2)
    add_text(
        doc,
        "The SCC-colony screening model uses normalized axial spacing, an exponential proximity term, and depth severity to derive an equivalent interacting crack depth and length. Collapse and fracture pressures are calculated separately, and the lower value governs. This interaction factor is an engineering screening heuristic, not a codified SCC interaction law.",
    )
    add_equation(doc, "s̄_i = s_i / √(Dt) ;     I = Σ exp(-0.85 s̄_i)")
    add_equation(doc, "F_int = 1 + min[0.65, 0.12I + 0.35(d_max/t)I] ;     d_eq = min(F_int d_max, 0.95t)")
    add_callout(
        doc,
        "Required advanced assessment",
        "SCC colonies near acceptance limits should be evaluated using a documented engineering critical assessment that addresses crack coalescence, residual stress, toughness, environment-assisted growth, pressure cycling, and inspection uncertainty.",
    )

    add_heading(doc, "7. Interacting-Anomaly Model", 1)
    add_text(
        doc,
        "Pairs are generated when axial edge spacing is within the selected interaction distance. Axial and circumferential spacing are normalized by √(Dt), then combined into a proximity measure. Mixed defect types, overlap, uncertainty, secondary stress, and residual stress increase the interaction penalty.",
    )
    add_equation(doc, "λ_x = s_x/√(Dt) ;     λ_θ = s_θ/√(Dt) ;     q = exp[-0.55(λ_x² + λ_θ²)]")
    add_equation(doc, "P_comb = min(P_1, P_2) / (F_int F_load)")
    add_text(
        doc,
        "Published FEA studies show that depth, length, orientation, overlap, and separation materially affect the pressure capacity of multiple corrosion defects [7]. Consequently, the interaction formula is used for prioritization and model generation; critical cases should be resolved with nonlinear FEA.",
    )

    add_page_break(doc)
    add_heading(doc, "8. Production FEA Specification", 1)
    add_heading(doc, "8.1 Mesh", 2)
    for item in (
        "Use quadratic solid or shell elements appropriate to the solver and defect thickness; ensure enough through-ligament resolution to recover strain gradients.",
        "Refine around MFL valleys, dent shoulders, crack fronts, crack tips, and interacting overlap zones; transition gradually to the remote mesh.",
        "Represent cracks with explicit free faces, cohesive/XFEM enrichment, or a validated fracture-mesh technique. A colored line without a discontinuity is not acceptable.",
        "Perform at least three mesh densities and demonstrate convergence of failure pressure, peak equivalent plastic strain away from singular crack tips, and fracture parameters.",
    ):
        add_bullet(doc, item)
    add_heading(doc, "8.2 Material and Loading", 2)
    add_table(
        doc,
        ["Model item", "Minimum requirement"],
        [
            ["Material", "True stress-true plastic strain curve; elastic modulus and Poisson ratio; grade-specific yield and tensile behavior"],
            ["Pressure", "Internal pressure on ID and credible internal flaw faces; ramp through the target or collapse load"],
            ["End load", "Closed-end axial pressure stress or explicit end-cap loading; avoid unintended axial restraint"],
            ["Secondary load", "Bending strain, thermal/soil load, dent constraint, and residual stress where applicable"],
            ["Failure criterion", "Validated plastic-collapse, local strain, damage, J/CTOD, FAD, or instability criterion"],
        ],
        [2200, 7160],
    )
    add_heading(doc, "8.3 Boundary Conditions", 2)
    add_text(
        doc,
        "A remote reference ring should suppress rigid-body motion while allowing realistic radial and axial deformation. Model length should be sufficient for boundary effects to decay before the anomaly region; the current framework uses 6D to 10D as a configurable starting range. Symmetry is permitted only when defect geometry, loading, and expected response are symmetric.",
    )
    add_heading(doc, "8.4 Solver Outputs", 2)
    for item in (
        "Pressure-displacement response and limit load.",
        "Equivalent plastic strain and remaining-ligament strain path.",
        "J-integral, CTOD, or stress-intensity measures for crack models.",
        "Contact and residual deformation for constrained dents.",
        "Fatigue-driving stress/strain ranges at stabilized cycles.",
    ):
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "9. Verification, Validation, and Uncertainty", 1)
    add_heading(doc, "9.1 Software Verification", 2)
    add_table(
        doc,
        ["Check", "Acceptance evidence"],
        [
            ["Parser", "Known TXT/XLSX fixtures produce identical normalized records"],
            ["Geometry", "Point coordinates, remaining ligament, spans, and removed crack elements match hand checks"],
            ["Screening", "B31G/RSTRENG/crack-growth benchmark calculations reproduce independent results"],
            ["Mesh", "No inverted elements, nonmanifold regions, unintended crack bridges, or overlapping labels"],
            ["Regression", "Automated unit, smoke, desktop, and mobile tests pass after each change"],
        ],
        [2200, 7160],
    )
    add_heading(doc, "9.2 Model Validation", 2)
    add_text(
        doc,
        "A production solver implementation should be validated in increasing order of complexity: pristine pipe closed-form hoop response; idealized single corrosion burst cases; measured-profile corrosion cases; smooth and constrained dent tests; single-crack fracture benchmarks; and interacting or SCC-colony cases. Validation should compare both global capacity and local response. The RSTRENG validation record and published realistic-profile FEA studies provide useful reference datasets [3, 8, 9].",
    )
    add_heading(doc, "9.3 Uncertainty Treatment", 2)
    for item in (
        "Apply tool-specific depth, length, width, and location tolerances; do not use one generic tolerance for all technologies.",
        "Preserve correlations and registration uncertainty when combining MFL, caliper, and crack-tool runs.",
        "Run nominal and conservative deterministic cases, then probabilistic cases where consequence and data quality warrant.",
        "Separate measurement uncertainty, model-form uncertainty, material variability, pressure-history uncertainty, and growth-rate uncertainty.",
    ):
        add_bullet(doc, item)
    add_callout(
        doc,
        "Decision rule",
        "The software output is a technical input to an integrity decision. Applicable regulation, operator procedures, qualified engineering review, and validation evidence govern operation, pressure reduction, excavation, repair, or replacement.",
        fill=PALE_BLUE,
    )

    add_heading(doc, "10. Maximum MOP and Reporting", 1)
    add_text(
        doc,
        "For each isolated feature and interacting pair, the workflow calculates a failure-pressure estimate and applies the assessment factor. The maximum recommended operating pressure is the minimum factored limit across all generated cases. The report identifies the governing feature or pair, failure mode, utilization, minimum fatigue life, source geometry, and raw-data point counts.",
    )
    add_equation(doc, "MOP_max = min_j (F_A P_failure,j) ;     utilization = MAOP / MOP_max")
    add_text(
        doc,
        "A reported maximum MOP should always include the assessment method, units, material basis, geometry source, uncertainty case, pressure-load definition, assessment factor, and whether the value came from screening or a validated nonlinear FEA solution.",
    )

    add_page_break(doc)
    add_heading(doc, "11. Implementation Limitations and Roadmap", 1)
    add_heading(doc, "Current Limitations", 2)
    for item in (
        "The 3D view is a geometric shell representation and does not currently solve the finite-element equilibrium equations.",
        "Raw MFL depth percent is treated as a wall-loss surface after interpolation; MFL signal inversion and probability-of-detection are outside the current scope.",
        "Caliper deformation does not yet reconstruct full ovality, local tool dynamics, or dent restraint from soil/rock contact.",
        "Crack element removal represents an open discontinuity but does not calculate a crack-tip singular field in the browser.",
        "SCC and mixed-defect interaction factors are screening heuristics requiring validation before operational acceptance decisions.",
        "Fatigue calculations use simplified Paris-law integration and do not yet include load-sequence, closure, overload retardation, or corrosion-fatigue effects.",
    ):
        add_bullet(doc, item)
    add_heading(doc, "Recommended Next Steps", 2)
    add_number(doc, "Add neutral mesh export (Abaqus INP and an open solver format) with node sets, element sets, crack faces, pressure surfaces, and material cards.")
    add_number(doc, "Add measurement registration, outlier handling, interpolation controls, and tool-specific uncertainty metadata.")
    add_number(doc, "Implement an external nonlinear solver adapter and import solved fields for contour visualization and report generation.")
    add_number(doc, "Create a benchmark library spanning corrosion, dents, cracks, SCC colonies, and mixed interacting anomalies.")
    add_number(doc, "Add review signatures, calculation revision control, and reproducible input/output packages for engineering governance.")

    add_heading(doc, "12. Conclusion", 1)
    add_text(
        doc,
        "The revised Automated ILI-to-FEA workflow closes the most important conceptual gap between inspection data and engineering assessment: the defect geometry shown to the engineer is the geometry used to size the calculation. Metal loss thins the shell, dents deform it, and cracks interrupt it. This creates a credible foundation for automated model generation while preserving a clear boundary between standards-based screening, engineering heuristics, and validated Level 3 nonlinear FEA.",
    )

    add_page_break(doc)
    add_heading(doc, "References", 1)
    references = [
        "[1] ASME. ASME B31G-2023, Manual for Determining the Remaining Strength of Corroded Pipelines. https://www.asme.org/codes-standards/find-codes-standards/b31g-manual-determining-remaining-strength-corroded-pipelines",
        "[2] API and ASME. API 579-1/ASME FFS-1, Fitness-For-Service. Overview: https://www.api.org/products-and-services/training/inspection-training",
        "[3] Kiefner, J. F., Vieth, P. H., and Roytman, I. Continued Validation of RSTRENG, Final Report, AGA 97003260, 1996. https://www.osti.gov/biblio/441672",
        "[4] Kiefner, J. F., and Vieth, P. H. PC Program Speeds New Criterion for Evaluating Corroded Pipe, Oil & Gas Journal 88(34), 1990. https://www.osti.gov/biblio/6828231",
        "[5] Paris, P. C., and Erdogan, F. A Critical Analysis of Crack Propagation Laws, Journal of Basic Engineering 85(4), 528-533, 1963. https://doi.org/10.1115/1.3656900",
        "[6] U.S. DOT PHMSA. Pipeline Corrosion, Final Report, 2008. https://www.phmsa.dot.gov/sites/phmsa.dot.gov/files/docs/technical-resources/pipeline/gas-distribution-integrity-management/65996/finalreportpipelinecorrosion.pdf",
        "[7] Ma, B. et al. Assessment by finite element modeling of the interaction of multiple corrosion defects and the effect on failure pressure of corroded pipelines, Engineering Structures 165, 278-286, 2018. https://doi.org/10.1016/j.engstruct.2018.03.040",
        "[8] Ferreira, A. D. M. et al. New procedure of automatic modeling of pipelines with realistic shaped corrosion defects, Engineering Structures 221, 111030, 2020. https://doi.org/10.1016/j.engstruct.2020.111030",
        "[9] Cabral, R. M. S. et al. Assessment by finite element modeling of pipelines with corrosion defects based on River-Bottom Profile model, Engineering Structures 261, 114246, 2022. https://doi.org/10.1016/j.engstruct.2022.114246",
        "[10] Automated finite element analysis of burst capacity for corroded pipelines, International Journal of Pressure Vessels and Piping 216, 105489, 2025. https://doi.org/10.1016/j.ijpvp.2025.105489",
        "[11] PHMSA. Dent Fatigue Life Assessment, Closeout Report. https://primis.phmsa.dot.gov/rd/FileGet/7287/DOT432_Closeout_report.pdf",
        "[12] PRCI. Understanding Magnetic Flux Leakage Signals from Mechanical Damage in Pipelines, 2009. https://www.prci.org/Research/KeyResults/18835.aspx",
    ]
    for reference in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        set_run_font(p.add_run(reference), 9.5, BLACK)

    add_heading(doc, "Appendix A. Implementation Traceability", 1)
    add_table(
        doc,
        ["Component", "Implementation location"],
        [
            ["Raw table import and schema normalization", "src/cepa_crossing/gui_server.py: rows_from_uploaded_table; parse_ili_raw_file"],
            ["ILI-to-FEA geometry and calculation coupling", "src/cepa_crossing/gui_server.py: calculate_ili_to_fea_payload"],
            ["3D defect-aware shell construction", "web/ili-fea-3d.js"],
            ["Optional feature/raw import interface", "web/index.html; web/app.js; web/styles.css"],
            ["Automated regression tests", "tests/test_gui_server.py; tests/gui_smoke.mjs"],
        ],
        [3100, 6260],
    )
    add_text(
        doc,
        "This appendix identifies the software locations implementing the methodology described in this paper. Line numbers are intentionally omitted because they change as the application evolves.",
        italic=True,
        color=MID_GRAY,
        after=0,
    )

    doc.core_properties.title = "Automated ILI-to-FEA Methodology White Paper"
    doc.core_properties.subject = "Pipeline defect-aware geometry, assessment, and FEA methodology"
    doc.core_properties.author = "Pipeline Engineering Assessment Software"
    doc.core_properties.keywords = "ILI, FEA, MFL, caliper, crack, SCC, pipeline, RSTRENG, B31G"
    doc.core_properties.comments = "Generated June 6, 2026."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
