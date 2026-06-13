from __future__ import annotations

import math
import sys
from pathlib import Path

import numpy as np
from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))
sys.path.insert(0, str(ROOT / "tools"))

from cepa_crossing.gui_server import (  # noqa: E402
    ILI_SURROGATE_BENCHMARKS,
    ILI_SURROGATE_MODELS,
    calculate_ili_to_fea_payload,
    ili_surrogate_predict,
    ili_surrogate_validation,
)
from build_ili_fea_validation_whitepaper import font, make_charts  # noqa: E402
from build_ili_fea_whitepaper import (  # noqa: E402
    BLACK,
    DARK_BLUE,
    GOLD,
    MID_GRAY,
    NAVY,
    PALE_BLUE,
    PALE_TEAL,
    add_bullet,
    add_callout,
    add_equation,
    add_header_footer,
    add_heading,
    add_number,
    add_rich_paragraph,
    add_table,
    add_text,
    configure_section,
    configure_styles,
    set_run_font,
)

OUTPUT_DIR = ROOT / "outputs"
ASSET_DIR = OUTPUT_DIR / "ili_fea_comprehensive_assets"
OUTPUT = OUTPUT_DIR / "Automated_ILI_to_FEA_Comprehensive_Validation_White_Paper.docx"

INK = "#203748"
GRID = "#8FB2C4"
TEAL = "#0F766E"
ORANGE = "#D97706"
PURPLE = "#7048A8"
RED = "#991B1B"
LIGHT = "#E4EEF3"


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_figure(doc: Document, path: Path, caption: str, width=6.35) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        doc_pr = drawing.find(".//wp:docPr", {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"})
        if doc_pr is not None:
            doc_pr.set("descr", caption)
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(10)
    set_run_font(caption_p.add_run(caption), 9.2, MID_GRAY, italic=True)


def workflow_chart() -> Path:
    path = ASSET_DIR / "calculation_workflow.png"
    image = Image.new("RGB", (1600, 1060), "white")
    draw = ImageDraw.Draw(image)
    draw.text((800, 38), "Automated ILI-to-FEA calculation workflow", fill=INK, font=font(32, True), anchor="ma")
    stages = [
        ("1", "Data ingestion", "Feature list or raw MFL, crack-tool, and caliper points"),
        ("2", "Geometry reconstruction", "Coordinates, wall loss, dent displacement, crack paths, SCC colonies"),
        ("3", "Single adaptive topology", "Remote cells transition to locally subdivided anomaly cells"),
        ("4", "Deterministic assessment", "B31G-family, fracture, dent, interaction, fatigue, strain, MOP"),
        ("5", "Rapid surrogate", "ANN or DNN normalized corrosion-capacity prediction"),
        ("6", "Reliability integration", "Monte Carlo demand-capacity propagation and reliability index"),
        ("7", "Engineering disposition", "Governing MOP, fatigue life, strain status, limitations, review"),
    ]
    left, width, height, gap = 170, 1260, 102, 28
    y = 105
    for index, (number, title, detail) in enumerate(stages):
        fill = "#E9F5F3" if index in {0, 1, 2} else "#EAF3F8" if index in {3, 4, 5} else "#F4F6F9"
        draw.rounded_rectangle((left, y, left + width, y + height), radius=8, fill=fill, outline="#7CA5B8", width=2)
        draw.rounded_rectangle((left + 18, y + 20, left + 80, y + 82), radius=6, fill=TEAL, outline=TEAL)
        draw.text((left + 49, y + 51), number, fill="white", font=font(24, True), anchor="mm")
        draw.text((left + 105, y + 28), title, fill=INK, font=font(23, True))
        draw.text((left + 105, y + 63), detail, fill="#475569", font=font(18))
        if index < len(stages) - 1:
            x = left + width // 2
            draw.line((x, y + height, x, y + height + gap - 7), fill="#54788B", width=4)
            draw.polygon([(x - 9, y + height + gap - 15), (x + 9, y + height + gap - 15), (x, y + height + gap)], fill="#54788B")
        y += height + gap
    image.save(path)
    return path


def project(x: float, theta: float, radial: float, width: int, height: int) -> tuple[float, float]:
    px = 165 + (x + 5.5) / 11.0 * (width - 330)
    py = height * 0.52 - math.sin(theta) * 185 * radial - math.cos(theta) * 30
    return px, py


def draw_model_reference(name: str, anomalies: list[dict], filename: str) -> Path:
    path = ASSET_DIR / filename
    width, height = 1500, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((width // 2, 32), name, fill=INK, font=font(30, True), anchor="ma")
    draw.text(
        (width // 2, 73),
        "Current raw-data reconstruction rules: one shell topology, local subdivision, anomaly-shaped surface changes",
        fill="#475569",
        font=font(18),
        anchor="ma",
    )

    local_x = []
    local_theta = []
    for anomaly in anomalies:
        local_x.extend(np.linspace(anomaly["x"] - anomaly["lx"], anomaly["x"] + anomaly["lx"], 15))
        local_theta.extend(np.linspace(anomaly["theta"] - anomaly["lt"], anomaly["theta"] + anomaly["lt"], 13))
    x_lines = sorted(set(np.linspace(-5.5, 5.5, 25).round(4).tolist() + np.array(local_x).round(4).tolist()))
    theta_lines = sorted(set(np.linspace(-math.pi, math.pi, 25).round(4).tolist() + np.array(local_theta).round(4).tolist()))

    def radial_state(x, theta):
        radial = 1.0
        color = GRID
        for anomaly in anomalies:
            dx = (x - anomaly["x"]) / max(anomaly["lx"], 0.1)
            dt = math.atan2(math.sin(theta - anomaly["theta"]), math.cos(theta - anomaly["theta"])) / max(anomaly["lt"], 0.04)
            weight = math.exp(-1.8 * (dx * dx + dt * dt))
            if anomaly["type"] == "metal_loss":
                radial -= anomaly.get("severity", 0.22) * weight
                if weight > 0.18:
                    color = ORANGE
            elif anomaly["type"] == "dent":
                radial -= anomaly.get("severity", 0.32) * weight
                if weight > 0.18:
                    color = PURPLE
        return radial, color

    crack_paths = [a for a in anomalies if a["type"] in {"crack", "scc"}]

    for theta in theta_lines:
        points = []
        colors = []
        for x in np.linspace(-5.5, 5.5, 260):
            missing = any(
                abs(x - a["x"]) < a["lx"]
                and abs(math.atan2(math.sin(theta - a["theta"]), math.cos(theta - a["theta"]))) < a["lt"] * 0.08
                for a in crack_paths
            )
            if missing:
                if len(points) > 1:
                    draw.line(points, fill=colors[-1], width=1)
                points = []
                continue
            radial, color = radial_state(x, theta)
            points.append(project(x, theta, radial, width, height))
            colors.append(color)
        if len(points) > 1:
            draw.line(points, fill=colors[-1] if colors else GRID, width=1)

    for x in x_lines:
        points = []
        for theta in np.linspace(-math.pi, math.pi, 220):
            missing = any(
                abs(x - a["x"]) < a["lx"]
                and abs(math.atan2(math.sin(theta - a["theta"]), math.cos(theta - a["theta"]))) < a["lt"] * 0.08
                for a in crack_paths
            )
            if missing:
                if len(points) > 1:
                    draw.line(points, fill=GRID, width=1)
                points = []
                continue
            radial, _ = radial_state(x, theta)
            points.append(project(x, theta, radial, width, height))
        if len(points) > 1:
            draw.line(points, fill=GRID, width=1)

    for anomaly in anomalies:
        if anomaly["type"] in {"metal_loss", "dent"}:
            anomaly_color = ORANGE if anomaly["type"] == "metal_loss" else PURPLE
            for theta_offset in np.linspace(-0.72 * anomaly["lt"], 0.72 * anomaly["lt"], 9):
                pts = []
                for x in np.linspace(anomaly["x"] - anomaly["lx"], anomaly["x"] + anomaly["lx"], 70):
                    radial, _ = radial_state(x, anomaly["theta"] + theta_offset)
                    pts.append(project(x, anomaly["theta"] + theta_offset, radial, width, height))
                draw.line(pts, fill=anomaly_color, width=2)
        if anomaly["type"] in {"crack", "scc"}:
            offsets = [0] if anomaly["type"] == "crack" else [-0.42, 0, 0.42]
            for offset in offsets:
                pts = []
                for x in np.linspace(anomaly["x"] - anomaly["lx"], anomaly["x"] + anomaly["lx"], 45):
                    theta = anomaly["theta"] + offset * anomaly["lt"] + 0.08 * anomaly["lt"] * math.sin((x - anomaly["x"]) * 8)
                    pts.append(project(x, theta, 1.008, width, height))
                draw.line(pts, fill=RED, width=6)
        px, py = project(anomaly["x"], anomaly["theta"], 1.12, width, height)
        label = anomaly.get("label", anomaly["type"].replace("_", " ").title())
        draw.rounded_rectangle((px - 78, py - 48, px + 78, py - 14), radius=5, fill="white", outline="#64748B")
        draw.text((px, py - 31), label, fill=INK, font=font(16, True), anchor="mm")

    draw.text((150, height - 54), "Remote mesh", fill=INK, font=font(18, True))
    draw.line((280, height - 45, 390, height - 45), fill=GRID, width=2)
    draw.text((500, height - 54), "Locally subdivided cells", fill=INK, font=font(18, True))
    for offset in range(0, 112, 16):
        draw.line((720 + offset, height - 58, 720 + offset, height - 32), fill="#315B70", width=1)
    draw.line((720, height - 58, 832, height - 58), fill="#315B70", width=1)
    draw.line((720, height - 32, 832, height - 32), fill="#315B70", width=1)
    draw.text((975, height - 54), "Crack paths remove intersected triangles", fill=INK, font=font(18, True))
    draw.line((1320, height - 45, 1410, height - 45), fill=RED, width=6)
    image.save(path)
    return path


def validation_evidence_chart() -> Path:
    path = ASSET_DIR / "published_validation_evidence.png"
    image = Image.new("RGB", (1550, 920), "white")
    draw = ImageDraw.Draw(image)
    draw.text((775, 35), "Published validation evidence relevant to the implemented workflow", fill=INK, font=font(29, True), anchor="ma")
    actual = np.array([row[4] for row in ILI_SURROGATE_BENCHMARKS])
    ann = np.array([ili_surrogate_predict("ann", *row[:4]) for row in ILI_SURROGATE_BENCHMARKS])
    dnn = np.array([ili_surrogate_predict("dnn", *row[:4]) for row in ILI_SURROGATE_BENCHMARKS])
    ann_error = (ann - actual) / actual * 100
    dnn_error = (dnn - actual) / actual * 100
    studies = [
        ("Current ANN reproduction", float(ann_error.min()), float(ann_error.max()), None, "[1]"),
        ("Current DNN reproduction", float(dnn_error.min()), float(dnn_error.max()), None, "[1]"),
        ("Longitudinal interaction ANN", -9.39, 4.63, None, "[1]"),
        ("Single-corrosion ANN", -6.86, 7.53, None, "[3]"),
        ("Plain-dent DNN vs experiments", None, None, 1.52, "[5]"),
    ]
    left, right, top, bottom = 430, 1450, 145, 775
    x_min, x_max = -12.0, 12.0
    x_map = lambda v: left + (v - x_min) / (x_max - x_min) * (right - left)
    for tick in range(-10, 11, 5):
        x = x_map(tick)
        draw.line((x, top, x, bottom), fill="#D9E2E8", width=1)
        draw.text((x, bottom + 18), f"{tick:+d}%", fill="#334155", font=font(18), anchor="ma")
    draw.line((x_map(0), top, x_map(0), bottom), fill="#334155", width=3)
    row_h = 112
    for i, (label, low, high, mape, citation) in enumerate(studies):
        y = top + 52 + i * row_h
        draw.text((30, y), f"{label} {citation}", fill=INK, font=font(19, True), anchor="lm")
        if low is not None:
            draw.line((x_map(low), y, x_map(high), y), fill=TEAL, width=10)
            draw.ellipse((x_map(low) - 7, y - 7, x_map(low) + 7, y + 7), fill=TEAL)
            draw.ellipse((x_map(high) - 7, y - 7, x_map(high) + 7, y + 7), fill=TEAL)
            close_range = x_map(high) - x_map(low) < 120
            draw.text((x_map(low), y - 20), f"{low:+.2f}%", fill="#475569", font=font(15), anchor="ms")
            draw.text((x_map(high), y + (27 if close_range else -20)), f"{high:+.2f}%", fill="#475569", font=font(15), anchor="ma" if close_range else "ms")
        else:
            draw.ellipse((x_map(mape) - 9, y - 9, x_map(mape) + 9, y + 9), fill=ORANGE)
            draw.text((x_map(mape) + 16, y), f"MAPE {mape:.2f}%", fill="#475569", font=font(17), anchor="lm")
    draw.text((940, 845), "Signed ranges are reported prediction errors; the dent value is reported MAPE and is not a signed range.", fill="#475569", font=font(17), anchor="ma")
    image.save(path)
    return path


def sample_payload() -> dict:
    return {
        "pipe": {
            "outside_diameter_mm": 762,
            "wall_thickness_mm": 9.5,
            "maop_mpa": 6.9,
            "smys_mpa": 448,
            "smts_mpa": 535,
            "elastic_modulus_mpa": 207000,
            "fracture_toughness_mpa_sqrt_m": 95,
            "assessment_factor": 0.72,
        },
        "loading": {
            "pressure_range_mpa": 1.5,
            "cycles_per_year": 1000,
            "applied_cycles": 10000,
            "bending_strain_percent": 0.2,
            "secondary_stress_mpa": 0,
            "residual_stress_fraction": 0,
            "paris_c": 1e-12,
            "paris_m": 3,
        },
        "model": {
            "geometry_source": "auto",
            "interaction_distance_mm": 500,
            "sizing_case": "conservative",
            "mesh_refinement": "standard",
            "solver_strategy": "implicit_riks",
            "screening_method": "modified_b31g",
            "class_location": "2",
            "surrogate_model": "ann",
            "reliability_samples": 2500,
            "depth_cov": 0.10,
            "pressure_cov": 0.03,
            "model_error_cov": 0.08,
            "strain_limit": 0.06,
        },
        "features": {
            "ids": ["F-101", "F-102", "F-103", "F-104"],
            "types": ["metal_loss", "crack", "dent", "metal_loss"],
            "depths_percent": [42, 58, 18, 64],
            "lengths_mm": [110, 45, 75, 180],
            "widths_mm": [70, 20, 90, 110],
            "clock_positions": ["3:00", "3:30", "12:00", "4:30"],
            "distances_m": [1250.0, 1250.2, 1280.2, 1315.6],
            "orientations_deg": [0, 10, 0, 0],
            "surfaces": ["external"] * 4,
            "reported_failure_pressures_mpa": [0, 8.1, 0, 6.8],
        },
        "raw_data": {
            "mfl_samples": [
                {"feature_id": "F-101", "axial_offset_mm": -40, "circumferential_offset_mm": -20, "depth_percent": 28},
                {"feature_id": "F-101", "axial_offset_mm": 0, "circumferential_offset_mm": 0, "depth_percent": 42},
                {"feature_id": "F-101", "axial_offset_mm": 38, "circumferential_offset_mm": 18, "depth_percent": 34},
                {"feature_id": "F-104", "axial_offset_mm": -60, "circumferential_offset_mm": -25, "depth_percent": 48},
                {"feature_id": "F-104", "axial_offset_mm": 0, "circumferential_offset_mm": 0, "depth_percent": 64},
                {"feature_id": "F-104", "axial_offset_mm": 65, "circumferential_offset_mm": 30, "depth_percent": 51},
            ],
            "crack_samples": [
                {"feature_id": "F-102", "axial_offset_mm": -20, "circumferential_offset_mm": -4, "depth_mm": 4.2, "opening_mm": 0.25, "crack_id": "C1"},
                {"feature_id": "F-102", "axial_offset_mm": 0, "circumferential_offset_mm": 0, "depth_mm": 5.51, "opening_mm": 0.42, "crack_id": "C1"},
                {"feature_id": "F-102", "axial_offset_mm": 20, "circumferential_offset_mm": 5, "depth_mm": 4.6, "opening_mm": 0.28, "crack_id": "C1"},
            ],
            "caliper_samples": [
                {"feature_id": "F-103", "axial_offset_mm": -30, "circumferential_offset_mm": -25, "radial_deformation_mm": -1.1},
                {"feature_id": "F-103", "axial_offset_mm": 0, "circumferential_offset_mm": 0, "radial_deformation_mm": -1.71},
                {"feature_id": "F-103", "axial_offset_mm": 30, "circumferential_offset_mm": 25, "radial_deformation_mm": -1.0},
            ],
        },
    }


def add_equation_block(doc: Document, title: str, equation: str, definitions: str, citation: str = "") -> None:
    add_heading(doc, title, 3)
    add_equation(doc, equation, definitions)
    if citation:
        add_text(doc, citation, italic=True, color=MID_GRAY, after=6)


def add_variable_tables(doc: Document) -> None:
    add_heading(doc, "Complete Model Variable Register", 1)
    add_text(doc, "The following tables list every user-controlled or raw-data variable consumed by the current Automated ILI-to-FEA calculation path. Units are mandatory unless identified as dimensionless.")
    add_heading(doc, "Pipe, Material, and Loading Variables", 2)
    add_table(
        doc,
        ["Symbol / field", "Units", "Meaning and use"],
        [
            ["D / outside_diameter_mm", "mm", "Pipe outside diameter; hoop-pressure conversion, normalization, spacing, mesh scale."],
            ["t / wall_thickness_mm", "mm", "Nominal wall; defect depth conversion, remaining ligament, pressure capacity, mesh scale."],
            ["P_MAOP / maop_mpa", "MPa", "Current pressure demand and MOP utilization."],
            ["SMYS", "MPa", "Yield strength; flow stress, axial-stress normalization, load modifiers."],
            ["SMTS", "MPa", "Tensile strength; flow stress cap and pristine pressure."],
            ["E", "MPa", "Elastic modulus; retained for mechanical definition and related modules."],
            ["K_mat", "MPa sqrt(m)", "Fracture toughness used in crack-pressure and K screening."],
            ["F_a", "-", "Assessment factor applied to failure pressure to obtain allowable pressure/MOP."],
            ["Delta P", "MPa", "Pressure-cycle range converted to hoop-stress range and Delta K."],
            ["N_y", "cycles/year", "Cycles per year for fatigue-life conversion."],
            ["N_applied", "cycles", "Applied cycles used for fatigue damage ratio."],
            ["epsilon_b", "%", "Existing bending strain added to dent/interaction strain screening."],
            ["sigma_sec", "MPa", "Secondary stress used in interaction load modifier and axial surrogate ratio."],
            ["f_res", "-", "Residual stress as a fraction of SMYS."],
            ["C, m", "Paris units, -", "Paris-law coefficient and exponent."],
        ],
        [2500, 1450, 5410],
    )
    add_heading(doc, "Feature and Raw ILI Variables", 2)
    add_table(
        doc,
        ["Input family", "Fields", "Role"],
        [
            ["Feature identity", "feature_id, type, surface", "Registration and anomaly model selection."],
            ["Feature location", "distance_m, clock_position, orientation_deg", "Axial coordinate, circumferential angle, crack orientation."],
            ["Feature size", "depth_percent, length_mm, width_mm", "Capacity, interaction, visualization footprint, refinement extent."],
            ["Reported capacity", "reported_failure_pressure_mpa", "Optional screening comparison."],
            ["Raw MFL", "axial_offset_mm, circumferential_offset_mm, depth_percent", "Local metal-loss depth field and effective geometry."],
            ["Raw crack tool", "axial_offset_mm, circumferential_offset_mm, depth_mm, opening_mm, crack_id, anomaly_type", "Crack front/path, opening, SCC colony grouping, mesh-triangle removal."],
            ["Raw caliper", "axial_offset_mm, circumferential_offset_mm, radial_deformation_mm", "Dent surface displacement and curvature representation."],
        ],
        [1900, 3800, 3660],
    )
    add_heading(doc, "Model-Control and Uncertainty Variables", 2)
    add_table(
        doc,
        ["Field", "Allowed/current meaning", "Calculation effect"],
        [
            ["geometry_source", "auto, raw, feature", "Selects raw-tool geometry or summarized feature fallback."],
            ["interaction_distance_mm", ">= 0", "Maximum axial edge spacing for pair-model creation."],
            ["sizing_case", "nominal, conservative, probabilistic", "Applies inspection tolerances and interaction multiplier."],
            ["mesh_refinement", "coarse, standard, fine", "Base grid and local subdivision: 3, 4, or 5."],
            ["solver_strategy", "implicit Riks, Newton-Raphson, explicit quasi-static", "Model-ready metadata; browser does not execute a nonlinear solver."],
            ["screening_method", "B31G-family, RSTRENG, CorLAS, SCC", "Primary isolated-feature method."],
            ["class_location", "1 to 4", "Risk ranking multiplier."],
            ["surrogate_model", "ANN or DNN", "Selects embedded corrosion response surface."],
            ["reliability_samples", "500 to 50,000", "Monte Carlo sample count."],
            ["depth_cov", "-", "Coefficient of variation for depth and scaled length sampling."],
            ["pressure_cov", "-", "Pressure-demand coefficient of variation."],
            ["model_error_cov", "-", "Multiplicative surrogate/model-form error."],
            ["strain_limit", "fraction", "B31.8-oriented screening threshold; default 0.06."],
            ["prediction_years, annual_growth_percent", "years, %/year", "Future depth and failure prediction in ILI ranking."],
        ],
        [2450, 3020, 3890],
    )


def add_equations(doc: Document) -> None:
    add_heading(doc, "Raw Equations Implemented in the Current Model", 1)
    add_callout(
        doc,
        "Scope",
        "These are the equations executed by the present software. Some are standards-derived screening relations, some are literature-based fracture/fatigue relations, and some are explicitly identified engineering heuristics for interaction, visualization, or advisory ranking. They must not be confused with a completed nonlinear finite-element solve.",
    )
    add_heading(doc, "Geometry Translation and Raw-Data Reconstruction", 2)
    add_equation_block(doc, "Clock-position conversion", "theta = 30 h + 0.5 m  [degrees]", "h and m are the clock hour and minute.")
    add_equation_block(doc, "Axial model coordinate", "x_i = 1000 (s_i - min(s))", "s_i is ILI chainage in metres; x_i is millimetres from the first feature.")
    add_equation_block(doc, "Depth conversion and ligament", "d_i = (d_i,% / 100)t ;  t_rem,i = max(t - d_i, 0.01t)", "d_i is defect depth and t_rem is remaining ligament.")
    add_equation_block(doc, "MFL surface interpolation used by the renderer", "w = exp[-1.7((Delta x/0.42L_v)^2 + (Delta theta/0.42Theta_v)^2)] ; Delta r_MFL = -0.42(d/t)w", "L_v and Theta_v are visualization footprints scaled from feature length and width. The coefficient is a rendering scale, not a physical radial displacement.")
    add_equation_block(doc, "Caliper surface interpolation used by the renderer", "w = exp[-1.45((Delta x/0.48L_v)^2 + (Delta theta/0.48Theta_v)^2)] ; Delta r_cal = 0.34 clamp(delta_cal/t,-1.5,1.5)w", "delta_cal is measured radial deformation. The normalized surface displacement is a visual reconstruction rule.")
    add_equation_block(doc, "Crack path distance and void rule", "d_seg = min ||q - [p_0 + clamp(((q-p_0).v)/(v.v),0,1)v]|| ; remove triangle if d_seg < tau_crack", "The centroid of a surface triangle is tested against each crack polyline. tau_crack is 0.064, 0.048, or 0.038 visualization units for coarse, standard, or fine rendering.")

    add_heading(doc, "Pressure Capacity and Interaction", 2)
    add_equation_block(doc, "Flow stress and pristine pressure", "sigma_flow = min(SMYS + 69 MPa, SMTS) ; P_flow = 2t sigma_flow / D ; P_0 = 2t SMTS / D", "P_flow supports screening; P_0 normalizes the ANN/DNN output.", "[1,3,7]")
    add_equation_block(doc, "Modified B31.G geometry parameter", "z = L^2/(Dt) ; M = sqrt(1 + 0.6275z - 0.003375z^2), z <= 50 ; M = 3.3 + 0.032z, z > 50", "M is the Folias bulging factor.", "[7,8]")
    add_equation_block(doc, "Modified B31.G failure stress and pressure", "sigma_f = sigma_flow [1 - 0.85(d/t)]/[1 - 0.85d/(tM)] ; P_f = 2t sigma_f/D", "The allowable pressure is F_a P_f.", "[7,8]")
    add_equation_block(doc, "Crack fracture screen", "sigma_fr = K_mat/[Y sqrt(pi a)] ; P_fr = sigma_fr / [(D/2t)cos^2(phi) + (D/4t)sin^2(phi)]", "Y = 1.12, a is crack depth in metres, and phi is crack orientation.")
    add_equation_block(doc, "Crack plastic-collapse cap", "P_pl = P_flow max[0.15, 1 - 0.35(d/t)] ; P_crack = min(P_fr, P_pl)", "This is the current simplified cap, not an API 579 FAD or full J-integral solution.")
    add_equation_block(doc, "Dent pressure screen", "P_dent = P_flow/[1 + 8(d/D) + 0.55(d/t)]", "Current engineering screening relation; detailed dents require API RP 1183-compatible assessment and validated FEA [5,9].")
    add_equation_block(doc, "Normalized anomaly spacing", "lambda_x = s_x/sqrt(Dt) ; lambda_theta = s_theta/sqrt(Dt) ; I_s = sqrt(lambda_x^2 + lambda_theta^2) ; q = exp(-0.55 I_s^2)", "s_x and s_theta are axial and circumferential edge spacing.")
    add_equation_block(doc, "Implemented interaction factor", "F_I = clamp{1 + q[0.08 + 0.55(d_1/t+d_2/t)/2]F_mix F_unc + F_overlap, 1, 2.5}", "F_mix increases for mixed, crack-metal-loss, and dent-crack pairs; F_unc reflects sizing treatment. This is an explicit software heuristic, not a code equation.")
    add_equation_block(doc, "Combined-load modifier and pair pressure", "F_L = 1 + 0.35|sigma_sec|/SMYS + 0.18f_res ; P_pair = min(P_f,1,P_f,2)/(F_I F_L)", "Pair allowable MOP is F_a P_pair.")

    add_heading(doc, "Fatigue, Strain, and Reliability", 2)
    add_equation_block(doc, "Pressure-to-stress and stress intensity", "Delta sigma = Delta P[(D/2t)cos^2(phi) + (D/4t)sin^2(phi)] ; Delta K = Y Delta sigma sqrt(pi a)", "Closed-end axial contribution is orientation-dependent.")
    add_equation_block(doc, "Paris-law growth and life", "da/dN = C(Delta K)^m ; N_rem approximately (a_crit-a_0)/[C(Delta K)^m]", "The interaction module uses the displayed constant-rate estimate; the standalone crack-growth module numerically integrates midpoint increments. Paris and Erdogan provide the classical basis [11].")
    add_equation_block(doc, "Interaction strain heuristic", "epsilon_eq = 0.002 + 0.08 max(d/t) + 0.035q + 0.01|sigma_sec|/SMYS", "This is a software screening heuristic and is not an FEA-computed equivalent plastic strain.")
    add_equation_block(doc, "Dent curvature strain", "kappa = 2d/(L/2)^2 + 2d/(W/2)^2 ; epsilon_dent = t kappa/2 ; epsilon_max = max(epsilon_b + epsilon_dent, epsilon_pair)", "A geometric bending-strain screen used for B31.8/API RP 1183 triage [9,10].")
    add_equation_block(doc, "Monte Carlo demand and capacity", "C_j = P_0 f_NN(x_j) max[N(1,COV_model),0.45] ; Q_j = P_MAOP max[N(1,COV_P),0.70]", "Depth and length are normally sampled and truncated to the surrogate domain.")
    add_equation_block(doc, "Probability of failure and reliability index", "P_f = count(Q_j >= C_j)/N ; beta = -Phi^-1[(n_f+0.5)/(N+1)]", "A continuity correction prevents infinite beta when zero failures occur.")
    add_equation_block(doc, "Governing operating pressure", "MOP_max = min(F_a P_isolated,i, F_a P_pair,k) ; U_MOP = P_MAOP/MOP_max", "The standards-based/interaction limit governs; the surrogate remains advisory.")

    add_heading(doc, "ANN and DNN Equations", 2)
    add_equation_block(doc, "Input normalization", "x_1 = 2[(d/t)-0.2]/0.6 - 1 ; x_2 = 2[(L/D)-0.2]/1.6 - 1 ; x_3 = 2(s/sqrt(Dt))/3 - 1 ; x_4 = 2(sigma_ax/SMYS) - 1", "Inputs map the embedded publication domain to approximately [-1,1].", "[1,2]")
    add_equation_block(doc, "Feed-forward propagation", "h^(l) = tanh(h^(l-1)W^(l) + b^(l)) ; y = h^(L-1)W^(L) + b^(L) ; f_NN = clamp(y,0.05,1.05)", "ANN architecture: 4-8-1. DNN architecture: 4-10-6-1.")
    add_equation_block(doc, "Physical pressure conversion and advisory lower bound", "P_NN = f_NN P_0 ; P_NN,lower = P_NN max[0.5, 1 - 1.645 COV_model]", "The advisory MOP is F_a P_NN,lower.")


def add_weight_tables(doc: Document) -> None:
    add_heading(doc, "Embedded Neural-Network Coefficients", 1)
    add_text(doc, "The following tables expose the numerical coefficients executed by the software. Rows are input or preceding-layer neurons; columns are receiving neurons. Biases are listed in the final row for each layer.")
    for model_name in ("ann", "dnn"):
        model = ILI_SURROGATE_MODELS[model_name]
        add_heading(doc, model["label"], 2)
        for layer_index, (weights, biases) in enumerate(zip(model["weights"], model["biases"]), 1):
            columns = len(biases)
            headers = ["From / bias"] + [f"N{j + 1}" for j in range(columns)]
            rows = []
            for i, row in enumerate(weights):
                rows.append([f"Input N{i + 1}"] + [f"{value:.8f}" for value in row])
            rows.append(["Bias"] + [f"{value:.8f}" for value in biases])
            width_first = 1500
            remaining = 7860
            widths = [width_first] + [remaining // columns] * columns
            widths[-1] += 9360 - sum(widths)
            add_heading(doc, f"Layer {layer_index}", 3)
            add_table(doc, headers, rows, widths)


def build_document(charts: dict[str, Path], assets: dict[str, Path]) -> None:
    result = calculate_ili_to_fea_payload(sample_payload())
    ann = ili_surrogate_validation("ann")
    dnn = ili_surrogate_validation("dnn")

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_header_footer(doc.sections[0])

    add_text(doc, "", after=72)
    add_text(doc, "COMPREHENSIVE TECHNICAL WHITE PAPER", bold=True, color=GOLD, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("Automated ILI-to-FEA"), 30, NAVY, bold=True)
    add_text(doc, "Equations, reconstruction workflow, mesh logic, validation evidence, uncertainty, and governance", color=DARK_BLUE, after=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Current implementation documented against published pipeline-integrity literature", italic=True, color=GOLD, after=84, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Prepared for the Pipeline Engineering Assessment Software", bold=True, color=NAVY, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Version 2.0 | June 11, 2026", color=MID_GRAY, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)
    add_heading(doc, "Executive Summary", 1)
    add_text(doc, "This white paper documents the complete calculation path currently implemented in the Automated ILI-to-FEA module. It includes raw equations, all model variables, geometry-reconstruction rules, adaptive mesh logic, deterministic capacity and fatigue calculations, ANN/DNN propagation, reliability integration, validation evidence, and explicit limitations.")
    add_callout(doc, "Critical distinction", "The browser creates a model-ready three-dimensional shell representation from ILI measurements, but it does not execute a commercial nonlinear finite-element solver. Capacity and risk outputs are produced by the equations documented here. The representation is therefore a geometry and meshing reference, not a solved stress/strain contour.")
    add_heading(doc, "Current Sample Output", 2)
    add_table(
        doc,
        ["Output", "Current sample result", "Interpretation"],
        [
            ["Translated features", str(result["outputs"]["feature_count"]), "Two metal-loss features, one crack, and one dent."],
            ["Interaction models", str(result["outputs"]["interaction_model_count"]), "Pairs created within the selected edge-spacing limit."],
            ["Governing MOP", f"{result['outputs']['maximum_mop_mpa']:.3f} MPa", f"Governed by {result['outputs']['governing_source']}."],
            ["Mesh estimate", f"{result['outputs']['raw_mesh_nodes']:,} nodes / {result['outputs']['raw_mesh_elements']:,} elements", "Model-preparation estimate, not a solver mesh file."],
            ["Local surface density", f"{result['outputs']['mesh_density_ratio']:.0f}x remote", "Standard 4 x 4 local subdivision gives 16 surface cells per remote cell."],
            ["Crack voids", f"{result['outputs']['removed_crack_elements']} removed elements", "Renderer removes triangles intersecting crack paths."],
            ["Reliability", f"Pf={result['outputs']['probability_of_failure']:.5f}; beta={result['outputs']['reliability_index_beta']:.3f}", "Advisory Monte Carlo result for the controlling corrosion surrogate."],
        ],
        [2200, 2650, 4510],
    )

    add_heading(doc, "Document Map", 2)
    for item in (
        "Sections 1-3 define scope, workflow, and all calculation variables.",
        "Sections 4-7 provide every implemented equation and explain adaptive mesh density.",
        "Sections 8-12 compare the model domain with published corrosion, dent, crack, SCC/fatigue, and RSTRENG evidence.",
        "Appendices provide neural-network weights, verification requirements, and references.",
    ):
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "1. Model Scope and Evidence Hierarchy", 1)
    add_text(doc, "Validation is separated into four evidence levels so that related literature is not overstated as direct verification.")
    add_table(
        doc,
        ["Level", "Meaning", "Evidence in this paper"],
        [
            ["A - Direct reproduction", "Same normalized variables and target response as the embedded surrogate.", "29 published FEA parameter cases from Lo et al. [1]."],
            ["B - Independent source-method validation", "Full-scale burst tests used to validate related nonlinear FEA/ANN methodologies.", "X52, X65, X80 corrosion studies [1-4]."],
            ["C - Adjacent defect validation", "Published FEA/experimental evidence for dents and dent-cracks.", "Plain-dent DNN and XFEM dent-crack studies [5,6]."],
            ["D - Standards and historical databases", "Code methods, experimental databases, and governance.", "RSTRENG, ASME B31.8, API RP 1183, Paris law [7-11]."],
        ],
        [1700, 3900, 3760],
    )
    add_text(doc, "Only Level A directly tests the embedded corrosion response surface. Levels B-D support the physical methods, model-development strategy, and assessment boundaries.")

    add_page_break(doc)
    add_heading(doc, "2. Calculation Workflow", 1)
    add_figure(doc, assets["workflow"], "Figure 1. End-to-end calculation workflow implemented by the Automated ILI-to-FEA module.")
    for step in (
        "Import and parse summarized ILI features and/or raw MFL, crack-tool, and caliper data.",
        "Register each sample to feature ID, chainage, clock position, and local axial/circumferential offsets.",
        "Translate measurements to model coordinates and update depth, length, width, crack grouping, and dent geometry.",
        "Generate one shell topology. Base cells are subdivided around anomaly footprints; crack-intersected triangles are omitted.",
        "Screen isolated defects, generate interacting pairs, estimate fatigue/crack growth, and determine the governing allowable pressure.",
        "Run the selected ANN/DNN for domain-limited rapid corrosion prediction and propagate uncertainty by Monte Carlo simulation.",
        "Combine pressure, fatigue, reliability, and strain outputs into an engineering disposition with stated limitations.",
    ):
        add_number(doc, step)

    add_page_break(doc)
    add_variable_tables(doc)

    add_page_break(doc)
    add_equations(doc)

    add_page_break(doc)
    add_heading(doc, "5. Variable Mesh Density and Geometry Coupling", 1)
    add_text(doc, "The renderer uses a single continuous shell geometry. It does not place an extra patch mesh on top of the pipe. Instead, each base surface cell is tested against the normalized footprint of every anomaly. The same base cell is subdivided before its triangles are created.")
    add_equation(doc, "rho = sqrt[(Delta x/L_v)^2 + (Delta theta/Theta_v)^2] ; n = n_local if rho <= 1 ; n = n_transition if 1 < rho <= 1.8 ; n = 1 otherwise", "n is the number of subdivisions along both cell directions.")
    add_table(
        doc,
        ["Refinement", "Base grid", "Transition subdivision", "Local subdivision", "Local surface-cell density"],
        [
            ["Coarse", "32 axial x 20 angular", "2 x 2", "3 x 3", "9x remote"],
            ["Standard", "48 axial x 30 angular", "2 x 2", "4 x 4", "16x remote"],
            ["Fine", "64 axial x 40 angular", "3 x 3", "5 x 5", "25x remote"],
        ],
        [1750, 2350, 1800, 1700, 1760],
    )
    add_text(doc, "A standard local cell therefore contains 4 x 4 = 16 surface subcells where the corresponding remote cell contains one. Transition cells contain 2 x 2 = 4 subcells, reducing the size jump between the remote and anomaly regions. Metal-loss and dent displacements are evaluated at every refined vertex, so additional vertices resolve steeper local gradients. Crack paths remove triangles from this same refined topology.")
    add_callout(doc, "Convergence requirement", "A finer visual topology is not evidence of numerical convergence. A production nonlinear FEA must repeat the solution with decreasing local element size and demonstrate stable failure pressure, relevant strain away from singularities, and fracture parameters [3,5,6,12].")
    add_figure(doc, assets["interacting"], "Figure 2. Current-model interacting metal-loss reference. The denser lines are subdivisions of the existing shell topology, not a superimposed patch.")

    add_page_break(doc)
    add_heading(doc, "6. Direct Surrogate Verification", 1)
    add_text(doc, "The embedded ANN and DNN were checked against 29 normalized FEA cases drawn from the parameter trends reported by Lo, Karuppanan, and Ovinis [1]. The four inputs are d/t, L/D, s/sqrt(Dt), and axial stress/SMYS; the target is normalized failure pressure.")
    add_figure(doc, charts["parity"], "Figure 3. Embedded ANN and DNN predictions against the 29 published corrosion-FEA cases [1].")
    add_table(
        doc,
        ["Metric", "ANN 4-8-1", "DNN 4-10-6-1"],
        [
            ["R-squared", f"{ann['r_squared']:.6f}", f"{dnn['r_squared']:.6f}"],
            ["MAE, normalized pressure", f"{ann['mae_normalized_pressure']:.6f}", f"{dnn['mae_normalized_pressure']:.6f}"],
            ["MAPE", f"{ann['mape_percent']:.3f}%", f"{dnn['mape_percent']:.3f}%"],
            ["Maximum absolute error", f"{ann['maximum_absolute_error_percent']:.3f}%", f"{dnn['maximum_absolute_error_percent']:.3f}%"],
        ],
        [3600, 2880, 2880],
    )
    add_figure(doc, charts["residual"], "Figure 4. Signed residuals across the 29 direct-reproduction cases.")
    add_figure(doc, charts["sweeps"], "Figure 5. Published and predicted depth, length, and axial-stress response trends [1].")
    add_figure(doc, assets["single_corrosion"], "Figure 6. Current-model raw MFL metal-loss representation used as the geometry reference for corrosion-surrogate verification.")
    add_text(doc, "These are reproduction results, not blind validation. The DNN's extremely small error must be interpreted in that context. Lo et al. reported R-squared 0.9921 and errors from -9.39% to +4.63% for unseen longitudinal-interaction cases, providing a stronger independent check of the source ANN method [1].")

    add_page_break(doc)
    add_heading(doc, "7. Additional Published Corrosion Validation", 1)
    add_figure(doc, assets["evidence"], "Figure 7. Reported validation error evidence. Values are not pooled because datasets, metrics, pipe grades, and defect configurations differ.")
    add_table(
        doc,
        ["Publication", "Configuration", "Reported validation evidence", "Relevance"],
        [
            ["Lo et al. 2021 [1]", "X65 longitudinally interacting corrosion plus axial compression", "Unseen-case R2 0.9921; error -9.39% to +4.63%.", "Direct source domain for the four-input interaction surrogate."],
            ["Vijaya Kumar et al. 2022 [2]", "X80 circumferentially aligned interacting corrosion", "Related ANN and full-scale burst comparisons.", "Supports interacting-corrosion FEA/ANN strategy; different interaction orientation."],
            ["Kumar et al. 2021 [3]", "X80-X100 single corrosion plus axial compression", "Validation error -6.86% to +7.53%; R2 about 0.997.", "Supports depth, length, and axial-stress variables."],
            ["Zhu 2025 [12]", "Thin- and thick-wall corrosion with large FEA and burst datasets", "Numerical and experimental validation across machined and real corrosion.", "Highlights D/t applicability and need for thick-wall-specific models."],
            ["Kiefner et al. 1996 [7]", "Real and machined metal-loss burst database", "Continued validation of RSTRENG effective-area assessment.", "Supports actual-profile methods and conservative cross-checking."],
        ],
        [1800, 2300, 2700, 2560],
    )
    add_figure(doc, charts["burst"], "Figure 8. Six published full-scale burst comparisons supporting the source nonlinear FEA methodology [2].")
    add_callout(doc, "Applicability boundary", "The embedded surrogate is limited to the stated normalized ranges and corrosion-type response. It is not validated for D/t below the source range, arbitrary real-corrosion morphology, cracks, dents, SCC, weld flaws, or mixed anomalies.")

    add_page_break(doc)
    add_heading(doc, "8. Dent Verification Evidence", 1)
    add_text(doc, "Oh et al. developed and tested a DNN for unconstrained plain dents using nonlinear FEA and experimental results. Reported FEA-comparison MAPEs ranged from 0.33% to 5.55% across X52, X65, and X80 datasets, while comparison with three experiments produced MAPE 1.52% [5]. This validates the general use of nonlinear FEA plus machine learning for plain-dent burst prediction, but it does not directly validate the current software's simplified dent pressure or curvature-strain equations.")
    add_figure(doc, assets["dent"], "Figure 9. Current-model raw caliper dent representation. Caliper displacement changes vertices of the same locally refined shell.")
    add_text(doc, "The current module therefore treats dent outputs as screening. API RP 1183 requires dent shape, constraint, fatigue, coincident features, material properties, pressure history, and uncertainty to be evaluated using a qualified procedure [9].")

    add_page_break(doc)
    add_heading(doc, "9. Crack, Dent-Crack, and SCC/Fatigue Evidence", 1)
    add_text(doc, "Okodi et al. calibrated and validated XFEM models against full-scale tests for restrained and unrestrained X70 dent-crack specimens [6]. Their study demonstrates that crack depth, crack length, dent depth, restraint, and denting pressure materially affect propagation and burst. Those variables exceed the simplified current crack-pressure and dent-interaction screens.")
    add_figure(doc, assets["dent_crack"], "Figure 10. Current-model dent-crack representation: caliper-driven shell displacement with triangles omitted along the crack path.")
    add_text(doc, "The current fatigue calculation follows the Paris-law form da/dN = C(Delta K)^m [11]. It does not model crack closure, threshold variability, retardation, variable-amplitude sequence effects, residual stress redistribution, or a changing geometry factor during growth. Detailed SCC and crack assessments require qualified fracture mechanics and representative toughness.")
    add_figure(doc, assets["scc"], "Figure 11. Current-model SCC colony representation. Multiple crack IDs become separate crack paths, each represented by omitted triangles and a crack-front line.")
    add_callout(doc, "Verification status", "The crack and SCC representations have geometry/data-flow verification only. Published XFEM and full-scale results support the need for explicit crack propagation models, but no claim is made that the browser void representation reproduces XFEM/J-integral results.")

    add_page_break(doc)
    add_heading(doc, "10. Reliability and Uncertainty Propagation", 1)
    add_text(doc, "Reliability is calculated separately from deterministic pressure margin. The present Monte Carlo implementation samples depth, scaled length, pressure demand, and multiplicative surrogate error. It reports empirical failure probability and a continuity-corrected reliability index.")
    add_table(
        doc,
        ["Uncertain quantity", "Current distribution", "Important missing refinements"],
        [
            ["Depth ratio", "Truncated normal using depth COV", "Tool-specific bias, POD, correlation with length/width, excavation calibration."],
            ["Length ratio", "Truncated normal; 0.6 times depth COV", "Independent length uncertainty and feature-registration error."],
            ["Pressure demand", "Lower-truncated normal multiplier", "Time history, transients, operating controls, pressure correlation."],
            ["Model error", "Lower-truncated normal multiplier", "Bias separated by defect class, pipe grade, and model domain."],
            ["Material and toughness", "Not sampled in current surrogate path", "Population distributions, spatial correlation, heat/lot data."],
            ["Growth", "Not propagated in reliability loop", "Corrosion/SCC growth uncertainty and inspection interval."],
        ],
        [2200, 3000, 4160],
    )
    add_text(doc, "A zero observed failure count does not imply zero risk. The implemented plotting probability (n_f + 0.5)/(N + 1) prevents an infinite reliability index, but rare-event decisions require more samples or specialized reliability methods.")

    add_page_break(doc)
    add_heading(doc, "11. Verification Matrix and Acceptance Criteria", 1)
    add_table(
        doc,
        ["Component", "Current evidence", "Required production acceptance"],
        [
            ["Raw-data parser", "Unit tests and editable imported data", "Tool-vendor schemas, malformed-file tests, unit and datum verification."],
            ["Geometry reconstruction", "Visual/data-flow checks and model-reference figures", "Laser scan or excavation comparison; depth, area, curvature, and registration error."],
            ["Adaptive topology", "Single-mesh subdivision and crack-void checks", "Element-quality metrics and solver-compatible export."],
            ["Corrosion surrogate", "29-case reproduction plus source publications", "Frozen independent blind holdout and operator-specific bias study."],
            ["Interaction heuristic", "Trend consistency with source literature", "Direct calibration against interacting full-scale tests or validated nonlinear FEA."],
            ["Dent screen", "Adjacent published FEA/DNN evidence", "API RP 1183-qualified method and full dent-profile validation."],
            ["Crack/SCC", "Paris/fracture equations and adjacent XFEM evidence", "FAD/J-integral/XFEM validation with toughness and crack morphology."],
            ["Reliability", "Monte Carlo implementation checks", "Calibrated distributions, correlations, target reliability, rare-event convergence."],
        ],
        [2000, 3300, 4060],
    )

    add_heading(doc, "12. Limitations and Decision Rules", 1)
    for item in (
        "The 3D model is a reconstruction and mesh-reference visualization, not a nonlinear solver result.",
        "The interaction factor, dent pressure relation, and equivalent-strain relation are current software screening heuristics and need direct calibration before safety-critical use.",
        "ANN/DNN predictions are advisory and restricted to the published normalized corrosion domain.",
        "Published validations from different pipe grades and defect classes are corroborating evidence, not interchangeable datasets.",
        "MOP must be governed by an applicable standard method or a qualified nonlinear FEA/ECA, including the correct code edition and operator procedures.",
        "Mesh density must be supported by convergence, element-quality, material-model, boundary-condition, and failure-criterion verification.",
        "Fatigue and SCC require pressure-history quality, environmental effects, toughness, residual stress, crack morphology, and growth-law calibration.",
    ):
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "13. Qualification Roadmap", 1)
    roadmap = [
        "Freeze versioned equations, coefficients, datasets, and expected outputs under model-change control.",
        "Create solver-ready shell/solid exports preserving raw MFL, caliper, and crack geometry.",
        "Run element-size and model-length convergence studies for every defect class and interaction type.",
        "Validate reconstructed geometry against excavation, laser scan, or high-resolution metrology.",
        "Build independent blind datasets for single corrosion, interacting corrosion, dents, cracks, SCC colonies, and mixed defects.",
        "Calibrate model bias and uncertainty by pipe grade, D/t, inspection technology, defect class, and loading regime.",
        "Compare deterministic and probabilistic outputs against operator-approved acceptance criteria and target reliability.",
        "Require qualified engineering review whenever a case is outside domain, near an acceptance boundary, or controlled by crack/dent interaction.",
    ]
    for item in roadmap:
        add_number(doc, item)

    add_heading(doc, "14. Conclusion", 1)
    add_text(doc, "The current module provides an integrated and transparent engineering workflow: raw ILI data are translated into anomaly geometry; one adaptive shell topology represents metal loss, dents, cracks, and SCC paths; deterministic screening, fatigue, surrogate, and reliability calculations are executed; and a governing pressure and disposition are reported. The corrosion surrogate reproduces its selected published FEA cases closely, and multiple primary publications support the broader FEA-plus-machine-learning approach. The evidence does not yet constitute full qualification of every defect class or of a nonlinear FEA solver. The equations, variables, limitations, and verification roadmap in this paper define exactly what is implemented and what remains to be validated.")

    add_page_break(doc)
    add_weight_tables(doc)

    add_page_break(doc)
    add_heading(doc, "References", 1)
    references = [
        "[1] Lo, M.; Karuppanan, S.; Ovinis, M. Failure Pressure Prediction of a Corroded Pipeline with Longitudinally Interacting Corrosion Defects Subjected to Combined Loadings Using FEM and ANN. Journal of Marine Science and Engineering 9, 281 (2021). https://doi.org/10.3390/jmse9030281",
        "[2] Vijaya Kumar, S.D.; Karuppanan, S.; Ovinis, M. Artificial Neural Network-Based Failure Pressure Prediction of API 5L X80 Pipeline with Circumferentially Aligned Interacting Corrosion Defects Subjected to Combined Loadings. Materials 15, 2259 (2022). https://doi.org/10.3390/ma15062259 ; https://pmc.ncbi.nlm.nih.gov/articles/PMC8953741/",
        "[3] Kumar, S.D.V.; Karuppanan, S.; Ovinis, M. Failure Pressure Prediction of High Toughness Pipeline with a Single Corrosion Defect Subjected to Combined Loadings Using ANN. Metals 11, 373 (2021). https://doi.org/10.3390/met11020373",
        "[4] Lo, M.; Karuppanan, S.; Ovinis, M. A Review of Finite Element Analysis and Artificial Neural Networks as Failure Pressure Prediction Tools for Corroded Pipelines. Materials 14 (2021). https://pmc.ncbi.nlm.nih.gov/articles/PMC8538846/",
        "[5] Oh, D.; et al. Burst Pressure Prediction of API 5L X-Grade Dented Pipelines Using a Deep Neural Network. Journal of Marine Science and Engineering 8, 766 (2020). https://doi.org/10.3390/jmse8100766",
        "[6] Okodi, A.; Li, Y.; Cheng, R.; Kainat, M.; Yoosef-Ghodsi, N.; Adeeb, S. Crack Propagation and Burst Pressure of Pipeline with Restrained and Unrestrained Concentric Dent-Crack Defects Using XFEM. Applied Sciences 10, 7554 (2020). https://doi.org/10.3390/app10217554",
        "[7] Kiefner, J.F.; Vieth, P.H.; Roytman, I. Continued Validation of RSTRENG, Final Report (1996). https://www.osti.gov/biblio/441672",
        "[8] ASME. ASME B31G, Manual for Determining the Remaining Strength of Corroded Pipelines. https://www.asme.org/codes-standards/find-codes-standards/b31g-manual-determining-remaining-strength-corroded-pipelines",
        "[9] American Petroleum Institute. API Recommended Practice 1183, Assessment and Management of Dents in Pipelines. https://www.api.org/products-and-services/standards/important-standards-announcements/rp1183",
        "[10] ASME. ASME B31.8, Gas Transmission and Distribution Piping Systems. https://www.asme.org/codes-standards/find-codes-standards/b31-8-gas-transmission-distribution-piping-systems",
        "[11] Paris, P.; Erdogan, F. A Critical Analysis of Crack Propagation Laws. Journal of Basic Engineering 85(4), 528-533 (1963). https://doi.org/10.1115/1.3656900",
        "[12] Zhu, X.-K. Burst Pressure Models and Validations for Thick-Walled Pipelines Containing Corrosion Defects. Journal of Pipeline Science and Engineering (2025). https://doi.org/10.1016/j.jpse.2025.100301 ; accepted manuscript: https://www.osti.gov/pages/biblio/3017239",
    ]
    for reference in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.line_spacing = 1.15
        set_run_font(p.add_run(reference), 9.3, BLACK)

    doc.core_properties.title = "Automated ILI-to-FEA Comprehensive Validation White Paper"
    doc.core_properties.subject = "Equations, variables, adaptive mesh, validation, and qualification of the Automated ILI-to-FEA workflow"
    doc.core_properties.author = "Pipeline Engineering Assessment Software"
    doc.core_properties.keywords = "ILI, FEA, equations, adaptive mesh, ANN, DNN, validation, corrosion, dent, crack, SCC, reliability"
    doc.save(OUTPUT)


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    charts = make_charts()
    assets = {
        "workflow": workflow_chart(),
        "evidence": validation_evidence_chart(),
        "single_corrosion": draw_model_reference(
            "Single metal-loss verification reference",
            [{"type": "metal_loss", "x": 0.0, "theta": 0.30, "lx": 1.25, "lt": 0.36, "severity": 0.28, "label": "MFL metal loss"}],
            "model_single_corrosion.png",
        ),
        "interacting": draw_model_reference(
            "Interacting corrosion verification reference",
            [
                {"type": "metal_loss", "x": -0.65, "theta": 0.25, "lx": 1.05, "lt": 0.30, "severity": 0.22, "label": "Metal loss A"},
                {"type": "metal_loss", "x": 0.75, "theta": 0.32, "lx": 1.15, "lt": 0.34, "severity": 0.30, "label": "Metal loss B"},
            ],
            "model_interacting_corrosion.png",
        ),
        "dent": draw_model_reference(
            "Plain-dent verification reference",
            [{"type": "dent", "x": 0.0, "theta": 1.10, "lx": 1.35, "lt": 0.48, "severity": 0.35, "label": "Caliper dent"}],
            "model_dent.png",
        ),
        "dent_crack": draw_model_reference(
            "Dent-crack verification reference",
            [
                {"type": "dent", "x": 0.0, "theta": 0.85, "lx": 1.45, "lt": 0.46, "severity": 0.32, "label": "Dent"},
                {"type": "crack", "x": 0.0, "theta": 0.85, "lx": 0.95, "lt": 0.30, "label": "Crack void"},
            ],
            "model_dent_crack.png",
        ),
        "scc": draw_model_reference(
            "SCC crack-colony verification reference",
            [{"type": "scc", "x": 0.0, "theta": 0.55, "lx": 1.5, "lt": 0.55, "label": "SCC colony"}],
            "model_scc_colony.png",
        ),
    }
    build_document(charts, assets)
    print(OUTPUT)


if __name__ == "__main__":
    main()
